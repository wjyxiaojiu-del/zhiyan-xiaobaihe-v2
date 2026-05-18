"""导出路由 - Protocol 导出为 MD/DOCX/Excel"""
import io
import re
import math
from flask import Blueprint, render_template, request, jsonify, send_file
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from data.protocol_meta import PROTOCOL_META
from config import PROTOCOL_DIR
import os

export_bp = Blueprint("export", __name__)


@export_bp.route("/export")
def export_page():
    return render_template("export.html", protocols=PROTOCOL_META)


@export_bp.route("/api/export/<pid>/<fmt>")
def export_protocol(pid, fmt):
    matching = next((p for p in PROTOCOL_META if p["id"] == pid), None)
    if not matching:
        return jsonify({"error": "Protocol not found"}), 404

    filepath = os.path.join(PROTOCOL_DIR, matching["file"])
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    if fmt == "md":
        md_content = _convert_to_markdown(content, matching)
        buf = io.BytesIO()
        buf.write(md_content.encode("utf-8"))
        buf.seek(0)
        return send_file(buf, mimetype="text/markdown",
                        as_attachment=True,
                        download_name=f"{pid}_{matching['name']}.md")

    elif fmt == "docx":
        doc = _create_docx(content, matching)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        as_attachment=True,
                        download_name=f"{pid}_{matching['name']}.docx")

    return jsonify({"error": "Unsupported format"}), 400


@export_bp.route("/api/export-excel", methods=["POST"])
def export_excel():
    data = request.json
    sheet_name = data.get("sheetName", "数据处理结果")
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    title = data.get("title", "")

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    font_header = Font(name="宋体", size=10.5, bold=True)
    font_data = Font(name="宋体", size=10.5)
    font_title = Font(name="宋体", size=14, bold=True)
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    start_row = 1
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        cell = ws.cell(row=1, column=1, value=title)
        cell.font = font_title
        cell.alignment = align_center
        start_row = 2

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=h)
        cell.font = font_header
        cell.alignment = align_center
        cell.border = thin_border
        cell.fill = PatternFill(start_color="D9E2C8", end_color="D9E2C8", fill_type="solid")

    for row_idx, row_data in enumerate(rows, start_row + 1):
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = font_data
            cell.alignment = align_center
            cell.border = thin_border

    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                val = str(cell.value)
                length = sum(2 if ord(c) > 127 else 1 for c in val)
                max_len = max(max_len, length)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)

    for row in ws.rows:
        max_lines = 1
        for cell in row:
            if cell.value:
                val = str(cell.value)
                col_width = ws.column_dimensions[cell.column_letter].width or 10
                chars_per_line = max(col_width / 2, 5)
                lines = math.ceil(len(val) / chars_per_line) if len(val) > chars_per_line else 1
                max_lines = max(max_lines, lines)
        ws.row_dimensions[row[0].row].height = max(max_lines * 15, 20)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"{sheet_name}.xlsx",
    )


# ========== 内部工具函数 ==========

def _set_run_font(run, size_pt=10.5, bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')


def _add_paragraph(doc, text, size=10.5, bold=False, align=None, space_after=Pt(3)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, size, bold)
    return p


def _add_checkbox_line(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("☐ ")
    _set_run_font(run, size)
    run = p.add_run(text)
    _set_run_font(run, size)
    return p


def _convert_to_markdown(content, meta):
    lines = content.split("\n")
    md = [f"# {meta['id']} {meta['name']}\n"]
    for line in lines:
        line = line.strip()
        if not line:
            md.append("")
            continue
        if re.match(r"^[一二三四五六七八九十]+[、.]", line):
            md.append(f"\n## {line}\n")
        elif re.match(r"^步骤\d+", line):
            md.append(f"\n### {line}\n")
        elif re.match(r"^\d+\.", line):
            md.append(f"- [ ] {line[2:]}")
        else:
            md.append(line)
    return "\n".join(md)


def _create_docx(content, meta):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(f"{meta['id']}  {meta['name']}")
    _set_run_font(run, 16, bold=True)

    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_p.paragraph_format.space_after = Pt(12)
    run = src_p.add_run("植研小白盒 · 标准化Protocol")
    _set_run_font(run, 9)
    run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("Protocol编号") or line.startswith("实验名称"):
            continue
        if re.match(r"^[一二三四五六七八九十]+[、.]", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '4',
                qn('w:space'): '1', qn('w:color'): '2E7D32',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            run = p.add_run(line)
            _set_run_font(run, 12, bold=True)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            continue
        if re.match(r"^步骤\d+", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("🔹 " + line)
            _set_run_font(run, 11, bold=True)
            continue
        if line.startswith("为什么：") or line.startswith("做对了：") or line.startswith("做错了："):
            _add_checkbox_line(doc, line)
            continue
        if re.match(r"^\d+\.", line):
            _add_checkbox_line(doc, line)
            continue
        _add_paragraph(doc, line)

    doc.add_paragraph()
    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(20)
    run = sign_p.add_run("实验日期：____________    操作人：____________    复核人：____________")
    _set_run_font(run, 10.5)
    return doc
