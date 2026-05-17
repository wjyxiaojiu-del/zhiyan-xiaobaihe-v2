"""
植研小白盒 v2 - Flask Web应用
功能：Protocol卡片展示 + 试剂计算器 + AI问答 + 仪器可视化指南
依赖安装：pip install flask anthropic langchain langchain-community chromadb sentence-transformers
运行方法：python app.py
"""

import os
import json
import math
import re
import sqlite3
import secrets
import threading
import base64
from datetime import datetime
from functools import wraps
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory, session, redirect, url_for, flash
from werkzeug.security import generate_password_hash, check_password_hash

# PDF解析
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("[WARN] pdfplumber未安装，PDF提取功能不可用。运行: pip install pdfplumber")

# MinerU 高质量PDF解析（支持表格、公式、多栏排版）
try:
    from mineru.cli.common import do_parse as mineru_do_parse
    HAS_MINERU = True
except ImportError:
    HAS_MINERU = False
    print("[WARN] mineru未安装，将使用pdfplumber作为PDF解析后端。运行: pip install mineru")

# Claude API
try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))

# ========== 路径配置 ==========
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROTOCOL_DIR = os.path.join(BASE_DIR, "protocol_docs")
INSTRUMENT_DIR = os.path.join(BASE_DIR, "instrument_guides")
DB_DIR = os.path.join(BASE_DIR, "chroma_db")
# Vercel 只读文件系统，数据库放 /tmp
USER_DB = os.path.join("/tmp", "users.db") if os.environ.get("VERCEL") else os.path.join(BASE_DIR, "users.db")


# ========== 用户数据库 ==========
def get_user_db():
    conn = sqlite3.connect(USER_DB)
    conn.row_factory = sqlite3.Row
    return conn


def init_user_db():
    conn = get_user_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT DEFAULT 'user',
            is_premium INTEGER DEFAULT 0,
            avatar TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_login TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS user_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            data_type TEXT NOT NULL,
            data_name TEXT NOT NULL,
            data_json TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS protocol_favorites (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            protocol_id TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(user_id, protocol_id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS protocol_ratings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            protocol_id TEXT NOT NULL,
            rating INTEGER NOT NULL,
            comment TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(user_id, protocol_id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS experiment_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            protocol_id TEXT DEFAULT '',
            content TEXT NOT NULL,
            tags TEXT DEFAULT '',
            status TEXT DEFAULT 'done',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS user_api_settings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL UNIQUE,
            provider TEXT DEFAULT 'anthropic',
            api_key TEXT DEFAULT '',
            model TEXT DEFAULT 'claude-sonnet-4-20250514',
            base_url TEXT DEFAULT '',
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS teams (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            owner_id INTEGER NOT NULL,
            invite_code TEXT UNIQUE NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (owner_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS team_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id INTEGER NOT NULL,
            user_id INTEGER NOT NULL,
            role TEXT DEFAULT 'member',
            joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(team_id, user_id),
            FOREIGN KEY (team_id) REFERENCES teams(id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS team_protocols (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id INTEGER NOT NULL,
            protocol_id TEXT NOT NULL,
            added_by INTEGER NOT NULL,
            is_private INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(team_id, protocol_id),
            FOREIGN KEY (team_id) REFERENCES teams(id)
        );
        CREATE TABLE IF NOT EXISTS extract_tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            task_type TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            file_name TEXT DEFAULT '',
            file_data TEXT DEFAULT '',
            structured_json TEXT DEFAULT '',
            raw_text TEXT DEFAULT '',
            issues_json TEXT DEFAULT '',
            result_json TEXT DEFAULT '',
            error_msg TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            completed_at TIMESTAMP
        );
    """)
    # 数据库索引
    for idx_sql in [
        "CREATE INDEX IF NOT EXISTS idx_user_data_uid ON user_data(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_favorites_uid ON protocol_favorites(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_favorites_pid ON protocol_favorites(protocol_id)",
        "CREATE INDEX IF NOT EXISTS idx_ratings_pid ON protocol_ratings(protocol_id)",
        "CREATE INDEX IF NOT EXISTS idx_ratings_uid ON protocol_ratings(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_logs_uid ON experiment_logs(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_tasks_uid ON extract_tasks(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_tasks_status ON extract_tasks(status)",
        "CREATE INDEX IF NOT EXISTS idx_teams_owner ON teams(owner_id)",
        "CREATE INDEX IF NOT EXISTS idx_teams_code ON teams(invite_code)",
        "CREATE INDEX IF NOT EXISTS idx_tm_team ON team_members(team_id)",
        "CREATE INDEX IF NOT EXISTS idx_tm_user ON team_members(user_id)",
    ]:
        conn.execute(idx_sql)
    # 创建默认管理员
    admin = conn.execute("SELECT id FROM users WHERE username='admin'").fetchone()
    if not admin:
        conn.execute(
            "INSERT INTO users (username, email, password_hash, role, is_premium) VALUES (?, ?, ?, ?, ?)",
            ("admin", "admin@zhiyan.com", generate_password_hash("admin123"), "admin", 1)
        )
    conn.commit()
    conn.close()


try:
    init_user_db()
except Exception as e:
    print(f"[WARN] 数据库初始化失败（Vercel环境可忽略）: {e}")


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        # 先检查session
        if "user_id" in session:
            return f(*args, **kwargs)
        # 再检查token（小程序用）
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
            uid = _verify_token(token)
            if uid:
                session["user_id"] = uid
                return f(*args, **kwargs)
        flash("请先登录", "warning")
        return redirect(url_for("login"))
    return decorated


def _generate_token(user_id):
    """生成简单token"""
    payload = f"{user_id}:{secrets.token_hex(16)}"
    return payload


def _verify_token(token):
    """验证token，返回user_id"""
    try:
        uid = int(token.split(":")[0])
        conn = get_user_db()
        row = conn.execute("SELECT id FROM users WHERE id=?", (uid,)).fetchone()
        conn.close()
        return uid if row else None
    except:
        return None


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        conn = get_user_db()
        user = conn.execute("SELECT role FROM users WHERE id=?", (session["user_id"],)).fetchone()
        conn.close()
        if not user or user["role"] != "admin":
            flash("需要管理员权限", "error")
            return redirect(url_for("home"))
        return f(*args, **kwargs)
    return decorated




@app.context_processor
def inject_user():
    from flask import g
    user = getattr(g, '_cached_user', None)
    if user is None and "user_id" in session:
        conn = get_user_db()
        user = conn.execute("SELECT id, username, email, role, avatar FROM users WHERE id=?",
                            (session["user_id"],)).fetchone()
        conn.close()
        g._cached_user = user
    return dict(current_user=user)

# ========== Protocol元数据（卡片展示用） ==========
PROTOCOL_META = [
    {"id": "P001", "name": "蒽酮比色法测可溶性糖含量", "category": "植物生理", "icon": "basic", "desc": "用蒽酮试剂和糖反应生成蓝绿色物质，颜色深浅代表糖含量多少，用分光光度计测吸光度就能算出来。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P001_蒽酮比色法测可溶性糖.txt", "color": "#2196F3"},
    {"id": "P002", "name": "考马斯亮蓝法", "category": "植物生理", "icon": "basic", "desc": "考马斯亮蓝染料和蛋白质结合后变蓝色，蛋白越多颜色越深，测595nm吸光度就能算蛋白含量。", "difficulty": 2, "tags": ['分光光度计', '离心机', '研钵'], "file": "P002_Bradford法测可溶性蛋白.txt", "color": "#FF9800"},
    {"id": "P003", "name": "TBA法测丙二醛", "category": "植物生理", "icon": "basic", "desc": "MDA是膜脂过氧化的产物，代表植物受伤害程度。MDA和TBA反应生成红色物质，测532nm吸光度就能算出来。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P003_TBA法测MDA含量.txt", "color": "#9C27B0"},
    {"id": "P004", "name": "过氧化物酶", "category": "植物生理", "icon": "basic", "desc": "POD能催化愈创木酚和H₂O₂反应生成茶褐色产物，470nm测颜色变化速率就能算出POD活性。", "difficulty": 2, "tags": ['分光光度计', '离心机', '研钵'], "file": "P004_POD活性测定.txt", "color": "#F44336"},
    {"id": "P005", "name": "过氧化氢酶", "category": "植物生理", "icon": "basic", "desc": "CAT能分解H₂O₂，通过测240nm处H₂O₂减少的速率就能算出CAT活性。", "difficulty": 2, "tags": ['分光光度计', '离心机', '研钵'], "file": "P005_CAT活性测定.txt", "color": "#00BCD4"},
    {"id": "P006", "name": "植物基因组DNA提取", "category": "分子生物", "icon": "basic", "desc": "用CTAB这种洗涤剂把植物细胞膜打破释放DNA，再用氯仿把蛋白和DNA分开，最后用异丙醇把DNA沉淀出来。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P006_植物基因组DNA提取.txt", "color": "#607D8B"},
    {"id": "P007", "name": "普通PCR反应体系与程序", "category": "分子生物", "icon": "basic", "desc": "PCR就像DNA的复印机，通过反复加热降温，把目标DNA片段从几个拷贝扩增到几百万个拷贝。", "difficulty": 2, "tags": ['离心机', '电泳槽'], "file": "P007_PCR反应体系.txt", "color": "#795548"},
    {"id": "P008", "name": "琼脂糖凝胶电泳", "category": "分子生物", "icon": "basic", "desc": "把DNA样品加到琼脂糖凝胶的小孔里，通电后DNA按大小分开，大片段跑得慢在上面，小片段跑得快在下面。", "difficulty": 3, "tags": ['电泳槽'], "file": "P008_琼脂糖凝胶电泳.txt", "color": "#E91E63"},
    {"id": "P009", "name": "总RNA提取", "category": "分子生物", "icon": "basic", "desc": "用TRIzol把细胞裂解，氯仿分层把RNA和蛋白/DNA分开，异丙醇沉淀出RNA，再用反转录酶把RNA变成cD", "difficulty": 2, "tags": ['离心机', '水浴锅', '研钵'], "file": "P009_总RNA提取与反转录.txt", "color": "#00BCD4"},
    {"id": "P010", "name": "液氮研磨与试剂母液配制", "category": "基础操作", "icon": "basic", "desc": "液氮研磨就是把植物组织冻到-196°C变脆，然后用研钵捣碎；母液就是高浓度的试剂原液，用的时候稀释就行。", "difficulty": 2, "tags": ['离心机', '研钵'], "file": "P010_液氮研磨与母液配制.txt", "color": "#FFC107"},
    {"id": "P011", "name": "过氧化物酶(POD)活性测定 - 愈创木酚法", "category": "植物生理", "icon": "basic", "desc": "POD催化H2O2把愈创木酚氧化成茶褐色，470nm测颜色变化速率就能算POD活性。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P011_过氧化物酶活性测定_李合生法.txt", "color": "#2196F3"},
    {"id": "P012", "name": "过氧化氢酶(CAT)活性测定 - 高锰酸钾滴定法", "category": "植物生理", "icon": "basic", "desc": "CAT分解H2O2，用高锰酸钾滴定剩余的H2O2，就能算出CAT分解了多少H2O2。", "difficulty": 2, "tags": ['离心机', '水浴锅', '研钵'], "file": "P012_过氧化氢酶活性测定_滴定法.txt", "color": "#FF9800"},
    {"id": "P013", "name": "丙二醛(MDA)含量测定 - TBA法", "category": "植物生理", "icon": "basic", "desc": "MDA是膜脂过氧化的产物，和TBA反应生成红色物质，测532nm算含量。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P013_丙二醛含量测定_李合生法.txt", "color": "#9C27B0"},
    {"id": "P014", "name": "可溶性糖含量测定 - 蒽酮比色法", "category": "植物生理", "icon": "basic", "desc": "糖在浓硫酸作用下脱水生成糠醛，和蒽酮反应生成蓝绿色，630nm测吸光度算糖含量。", "difficulty": 2, "tags": ['分光光度计', '水浴锅'], "file": "P014_可溶性糖含量测定_蒽酮法_李合生.txt", "color": "#F44336"},
    {"id": "P015", "name": "叶绿素含量测定 - 分光光度法", "category": "植物生理", "icon": "basic", "desc": "用乙醇提取叶绿素，测665nm和649nm吸光度，用公式算叶绿素a、b和类胡萝卜素含量。", "difficulty": 2, "tags": ['分光光度计', '研钵'], "file": "P015_叶绿素含量测定_分光光度法.txt", "color": "#00BCD4"},
    {"id": "P016", "name": "可溶性蛋白质含量测定 - 考马斯亮蓝G-250法", "category": "植物生理", "icon": "basic", "desc": "考马斯亮蓝和蛋白质结合变蓝，595nm测吸光度就能算蛋白质含量。", "difficulty": 2, "tags": ['分光光度计', '离心机', '研钵'], "file": "P016_可溶性蛋白质含量测定_考马斯亮蓝法.txt", "color": "#795548"},
    {"id": "P017", "name": "呼吸速率测定 - 小篮子法/广口瓶法", "category": "植物生理", "icon": "basic", "desc": "植物释放的CO2被Ba(OH)2吸收，用草酸滴定剩余碱液，算出CO2释放量就是呼吸速率。", "difficulty": 2, "tags": ['滴定管'], "file": "P017_呼吸速率测定_小篮子法.txt", "color": "#607D8B"},
    {"id": "P018", "name": "植物组织DNA提取与测定 - 盐溶法", "category": "分子生物", "icon": "basic", "desc": "用SDS裂解细胞释放DNA，氯仿-异戊醇去蛋白，乙醇沉淀DNA，二苯胺法测含量。", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P018_DNA提取与测定_盐溶法.txt", "color": "#607D8B"},
    {"id": "P019", "name": "植物组织中自由水和束缚水含量的测定", "category": "植物生理", "icon": "basic", "desc": "自由水未被细胞原生质胶体颗粒吸附而可以自由移动、蒸发和结冰，也 可 以作为溶剂", "difficulty": 1, "tags": ['基础器材'], "file": "P019_自由水和束缚水含量测定.txt", "color": "#FFC107"},
    {"id": "P020", "name": "植物组织水势的测定", "category": "植物生理", "icon": "basic", "desc": "植物组织水势测定", "difficulty": 1, "tags": ['基础器材'], "file": "P020_植物组织水势测定.txt", "color": "#4CAF50"},
    {"id": "P021", "name": "植物细胞渗透势的测定(质壁分离法)", "category": "植物生理", "icon": "basic", "desc": "渗透势测定_质壁分离法", "difficulty": 1, "tags": ['基础器材'], "file": "P021_渗透势测定_质壁分离法.txt", "color": "#2196F3"},
    {"id": "P022", "name": "钾离子对气孔开度的影响", "category": "植物生理", "icon": "basic", "desc": "钾离子对气孔开度的影响", "difficulty": 1, "tags": ['基础器材'], "file": "P022_钾离子对气孔开度的影响.txt", "color": "#FF9800"},
    {"id": "P023", "name": "植物伤流液中糖和氨基酸的鉴定", "category": "植物生理", "icon": "basic", "desc": "用蒽酮试剂处理可以鉴定伤流液中糖的存在，并可测定其含量", "difficulty": 1, "tags": ['分光光度计', '水浴锅'], "file": "P023_伤流液中糖和氨基酸鉴定.txt", "color": "#9C27B0"},
    {"id": "P024", "name": "植物根系活力的测定(TTC法)", "category": "植物生理", "icon": "basic", "desc": "根系活力测定_TTC法", "difficulty": 2, "tags": ['分光光度计', '研钵'], "file": "P024_根系活力测定_TTC法.txt", "color": "#F44336"},
    {"id": "P025", "name": "植物组织中金属元素的测定(原子吸收分光光度法)", "category": "植物生理", "icon": "basic", "desc": "金属元素测定_原子吸收法", "difficulty": 3, "tags": ['分光光度计'], "file": "P025_金属元素测定_原子吸收法.txt", "color": "#00BCD4"},
    {"id": "P026", "name": "植物体内硝态氮含量的测定", "category": "植物生理", "icon": "basic", "desc": "硝态氮含量测定", "difficulty": 2, "tags": ['分光光度计', '水浴锅'], "file": "P026_硝态氮含量测定.txt", "color": "#795548"},
    {"id": "P027", "name": "植物体内硝酸还原酶活力的测定", "category": "植物生理", "icon": "basic", "desc": "硝酸还原酶活力测定", "difficulty": 2, "tags": ['分光光度计'], "file": "P027_硝酸还原酶活力测定.txt", "color": "#607D8B"},
    {"id": "P028", "name": "用真空渗入法测定环境因子", "category": "植物生理", "icon": "basic", "desc": "真空渗人法可使叶肉细胞间隙充满水分而下沉", "difficulty": 1, "tags": ['水浴锅'], "file": "P028_真空渗入法测定环境因子.txt", "color": "#E91E63"},
    {"id": "P029", "name": "叶绿体色素提取、分离和理化性质", "category": "植物生理", "icon": "basic", "desc": "叶绿体色素提取分离和理化性质", "difficulty": 1, "tags": ['水浴锅'], "file": "P029_叶绿体色素提取分离和理化性质.txt", "color": "#FFC107"},
    {"id": "P030", "name": "希尔反应的观察", "category": "植物生理", "icon": "basic", "desc": "希尔反应的观察", "difficulty": 1, "tags": ['研钵'], "file": "P030_希尔反应的观察.txt", "color": "#4CAF50"},
    {"id": "P031", "name": "RuBP羧化酶(RuBPCO)活性测定", "category": "植物生理", "icon": "basic", "desc": "RuBP羧化酶活性测定", "difficulty": 2, "tags": ['离心机', '水浴锅'], "file": "P031_RuBP羧化酶活性测定.txt", "color": "#2196F3"},
    {"id": "P032", "name": "乙醇酸氧化酶活性测定", "category": "植物生理", "icon": "basic", "desc": "乙醇酸氧化酶活性测定", "difficulty": 2, "tags": ['分光光度计', '离心机'], "file": "P032_乙醇酸氧化酶活性测定.txt", "color": "#FF9800"},
    {"id": "P033", "name": "红外CO₂分析仪测定光合与呼吸速率", "category": "植物生理", "icon": "basic", "desc": "红外CO2分析仪测光合呼吸速率", "difficulty": 3, "tags": ['基础器材'], "file": "P033_红外CO2分析仪测光合呼吸速率.txt", "color": "#9C27B0"},
    {"id": "P034", "name": "氧电极法测定光合与呼吸速率", "category": "植物生理", "icon": "basic", "desc": "氧电极法测光合呼吸速率", "difficulty": 3, "tags": ['水浴锅'], "file": "P034_氧电极法测光合呼吸速率.txt", "color": "#F44336"},
    {"id": "P035", "name": "叶绿体光诱导荧光强度的测定", "category": "植物生理", "icon": "basic", "desc": "叶绿体光诱导荧光测定", "difficulty": 2, "tags": ['分光光度计', '离心机'], "file": "P035_叶绿体光诱导荧光测定.txt", "color": "#00BCD4"},
    {"id": "P036", "name": "叶绿体甘油醛-3-磷酸脱氢酶活性测定", "category": "植物生理", "icon": "basic", "desc": "甘油醛_3_磷酸脱氢酶活性测定", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P036_甘油醛_3_磷酸脱氢酶活性测定.txt", "color": "#795548"},
    {"id": "P037", "name": "微量定容测压法测定种子的呼吸速率", "category": "植物生理", "icon": "basic", "desc": "微量定容测压法测呼吸速率", "difficulty": 2, "tags": ['基础器材'], "file": "P037_微量定容测压法测呼吸速率.txt", "color": "#607D8B"},
    {"id": "P038", "name": "NBT法测定SOD活力", "category": "植物生理", "icon": "basic", "desc": "本实验依据超氧物歧化酶抑制氮蓝四唑( NBT) 在光下的还原作用来确定酶 活性大小", "difficulty": 2, "tags": ['分光光度计', '离心机', '研钵'], "file": "P038_氮蓝四唑NBT法测定SOD活力.txt", "color": "#E91E63"},
    {"id": "P039", "name": "淀粉酶活性的测定", "category": "植物生理", "icon": "basic", "desc": "β- 淀粉酶每次从淀粉的非还原端切下 一 分子麦芽  糖，又被称为糖化酶", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P039_淀粉酶活性的测定.txt", "color": "#FFC107"},
    {"id": "P040", "name": "脲酶活性的测定", "category": "植物生理", "icon": "basic", "desc": "脲酶活性的测定", "difficulty": 2, "tags": ['分光光度计', '水浴锅'], "file": "P040_脲酶K.txt", "color": "#4CAF50"},
    {"id": "P041", "name": "POD同工酶凝胶圆盘电泳", "category": "分子生物", "icon": "basic", "desc": "POD同工酶凝胶圆盘电泳", "difficulty": 3, "tags": ['离心机', '电泳槽'], "file": "P041_植物过氧化物酶同工酶的测定_凝胶圆盘电泳.txt", "color": "#00BCD4"},
    {"id": "P042", "name": "SDS-PAGE测定蛋白质分子量", "category": "分子生物", "icon": "basic", "desc": "8蛋白质相对分子质量的测定(SDS-聚丙烯酰胺凝胶电泳法)", "difficulty": 3, "tags": ['电泳槽'], "file": "P042_蛋白质相对分子质量的测定_SDS_聚丙烯酰胺凝胶电泳法.txt", "color": "#607D8B"},
    {"id": "P043", "name": "凯氏定氮法测总氮和蛋白氮", "category": "植物生理", "icon": "basic", "desc": "凯氏定氮法测总氮和蛋白氮", "difficulty": 2, "tags": ['滴定管'], "file": "P043_植物组织中总氮_蛋白氮含量的测定_微量凯氏法.txt", "color": "#9C27B0"},
    {"id": "P044", "name": "茚三酮法测游离氨基酸总量", "category": "植物生理", "icon": "basic", "desc": "氨基酸与茚三酮共热时，能定量地生成二酮茚胺", "difficulty": 2, "tags": ['分光光度计', '水浴锅', '研钵'], "file": "P044_植物组织中游离氨基酸总量的测定_茚三酮显色法.txt", "color": "#F44336"},
    {"id": "P045", "name": "谷物淀粉含量的测定(旋光法)", "category": "植物生理", "icon": "basic", "desc": "谷物淀粉含量的测定(旋光法)", "difficulty": 2, "tags": ['离心机', '水浴锅'], "file": "P045_谷物淀粉含量的测定_旋光法.txt", "color": "#00BCD4"},
    {"id": "P046", "name": "植物种子生命力的快速测定", "category": "植物生理", "icon": "basic", "desc": "这些种子在衰老死亡时，内含 荧光物质虽然没有改变，但由于生命力衰退或已经死亡的细胞原生质之透性增 加，当浸泡种", "difficulty": 1, "tags": ['基础器材'], "file": "P046_植物种子生命力的快速测定.txt", "color": "#795548"},
    {"id": "P047", "name": "植物组织中纤维素含量的测定", "category": "植物生理", "icon": "basic", "desc": "植物组织中纤维素含量的测定", "difficulty": 2, "tags": ['分光光度计', '水浴锅'], "file": "P047_植物组织中纤维素含量的测定.txt", "color": "#607D8B"},
    {"id": "P048", "name": "苯丙氨酸解氨酶(PAL)活性测定", "category": "植物生理", "icon": "basic", "desc": "苯丙氨酸解氨酶(PAL)活性测定", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P048_苯丙氨酸解氨酶PALase活性的测定.txt", "color": "#E91E63"},
    {"id": "P049", "name": "DNA琼脂糖凝胶电泳", "category": "分子生物", "icon": "basic", "desc": "DNA琼脂糖凝胶电泳", "difficulty": 3, "tags": ['分光光度计', '电泳槽'], "file": "P049_DNA的琼脂糖凝胶电泳.txt", "color": "#00BCD4"},
    {"id": "P050", "name": "RNA的聚丙烯酰胺凝胶电泳", "category": "分子生物", "icon": "basic", "desc": "聚丙烯酰胺凝胶电泳是以聚丙烯酰胺凝胶为载体进行电泳的方法(详见实 验28)", "difficulty": 3, "tags": ['电泳槽'], "file": "P050_RNA的聚丙烯酰胺凝胶电泳.txt", "color": "#607D8B"},
    {"id": "P051", "name": "植物组织ATP酶活性的测定", "category": "植物生理", "icon": "basic", "desc": "它存在于生物细 胞的多个部位，比如细胞质膜上、叶绿体类囊体膜上，对整个生命的维持有着重 要的作用", "difficulty": 3, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P051_植物组织ATP酶活性的测定.txt", "color": "#2196F3"},
    {"id": "P052", "name": "植物种子中主要不饱和脂肪酸的分离(反相纸层析法)", "category": "植物生理", "icon": "basic", "desc": "植物种子中主要不饱和脂肪酸的分离(反相纸层析法)", "difficulty": 3, "tags": ['基础器材'], "file": "P052_植物种子中主要不饱和脂肪酸的分离_反相纸层析法.txt", "color": "#FF9800"},
    {"id": "P053", "name": "种子粗脂肪含量的测定", "category": "植物生理", "icon": "basic", "desc": "脂 肪( fat) 广泛存在于油料植物种子和果实中，测定脂肪的含量，可以鉴别其 品质的优劣，也是油料作物选种和", "difficulty": 2, "tags": ['水浴锅', '研钵'], "file": "P053_种子粗脂肪含量的测定.txt", "color": "#9C27B0"},
    {"id": "P054", "name": "气相色谱法测定植物样品膜脂中脂肪酸的含量", "category": "植物生理", "icon": "basic", "desc": "高等植物中的膜脂主要是各种类脂，可以用氯仿一甲醇溶液研磨提取，在碱 性条件下水解出高级脂肪酸并制成甲酯后，即可", "difficulty": 3, "tags": ['离心机'], "file": "P054_气相色谱法测定膜脂中脂肪酸的含量.txt", "color": "#F44336"},
    {"id": "P055", "name": "气相色谱法测定乙烯含量", "category": "植物生理", "icon": "basic", "desc": "气相色谱法测定乙烯含量", "difficulty": 3, "tags": ['基础器材'], "file": "P055_气相色谱法测定乙烯含量.txt", "color": "#00BCD4"},
    {"id": "P056", "name": "酶联免疫吸附检测法(ELISA)测定植物激素含量", "category": "植物生理", "icon": "basic", "desc": "在 ELISA 中，抗原抗体反应的检测依靠酶标记物来实现，常用的酶有辣根过氧化物 酶和碱性磷酸酯酶", "difficulty": 3, "tags": ['分光光度计', '离心机', '研钵'], "file": "P056_酶联免疫吸附检测法ELISA测定植物激素含量.txt", "color": "#795548"},
    {"id": "P057", "name": "ABA和GA的分离与测定", "category": "植物生理", "icon": "basic", "desc": "再 对 纯 化 的ABA 和 GA 进行生物学鉴定或物理化学鉴定", "difficulty": 3, "tags": ['基础器材'], "file": "P057_植物体内脱落酸_赤霉素的分离和测定.txt", "color": "#607D8B"},
    {"id": "P058", "name": "赤霉素对α一淀粉酶的诱导形成", "category": "植物生理", "icon": "basic", "desc": "淀粉性种子在萌动过程中，胚释放出来的赤霉素能诱导糊粉层细胞中 a-    淀粉酶基 因 的 表 达，引起α- ", "difficulty": 2, "tags": ['分光光度计', '水浴锅'], "file": "P058_赤霉素对α一淀粉酶的诱导形成.txt", "color": "#E91E63"},
    {"id": "P059", "name": "植物激素对愈伤组织的形成和分化的影响", "category": "植物生理", "icon": "basic", "desc": "愈伤组织在适当培养条件下分化根和芽的现象称为再分化", "difficulty": 3, "tags": ['水浴锅'], "file": "P059_植物激素对愈伤组织的形成和分化的影响.txt", "color": "#FFC107"},
    {"id": "P060", "name": "类似生长素对种子萌发的影响", "category": "植物生理", "icon": "basic", "desc": "类似生长素对种子萌发的影响", "difficulty": 2, "tags": ['基础器材'], "file": "P060_类似生长素对种子萌发的影响.txt", "color": "#4CAF50"},
    {"id": "P061", "name": "植物春化和光周期现象的观察", "category": "植物生理", "icon": "basic", "desc": "植物春化和光周期现象的观察", "difficulty": 1, "tags": ['基础器材'], "file": "P061_植物春化和光周期现象的观察.txt", "color": "#2196F3"},
    {"id": "P062", "name": "抗坏血酸(维生素C)含量的测定", "category": "植物生理", "icon": "basic", "desc": "因此当用蓝色的碱性2,6- 二氯酚靛酚溶液滴定含有抗坏血酸的草酸溶液  时，其中的抗坏血酸可以将2,6 - 二", "difficulty": 1, "tags": ['分光光度计', '研钵', '滴定管'], "file": "P062_抗坏血酸_维生素C_含量的测定.txt", "color": "#FF9800"},
    {"id": "P063", "name": "谷类作物种子中赖氨酸含量的测定", "category": "植物生理", "icon": "basic", "desc": "谷类作物种子中赖氨酸含量的测定", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P063_谷类作物种子中赖氨酸含量的测定.txt", "color": "#9C27B0"},
    {"id": "P064", "name": "脯氨酸含量的测定", "category": "植物生理", "icon": "basic", "desc": "脯氨酸含量的测定", "difficulty": 2, "tags": ['分光光度计', '离心机', '水浴锅'], "file": "P064_脯氨酸含量的测定.txt", "color": "#F44336"},
    {"id": "P065", "name": "电导仪法测定植物抗逆性", "category": "植物生理", "icon": "basic", "desc": "植物细胞膜对维持细胞的微环境和正常的代谢起着重要的作用", "difficulty": 2, "tags": ['水浴锅'], "file": "P065_植物抗逆性的测定_电导仪法.txt", "color": "#00BCD4"}
]

# ========== 仪器元数据 ==========
INSTRUMENT_META = [
    {
        "id": "I001", "name": "高速冷冻离心机",
        "icon": "centrifuge",
        "desc": "低温高速旋转分离液体中的固体和液体",
        "protocols": ["P002", "P003", "P004", "P005", "P006", "P011", "P012", "P013", "P016", "P018"],
        "color": "#1565C0",
        "hotspots": [
            {"name": "控制面板", "desc": "设置温度、转速、时间", "x": "30%", "y": "20%", "w": "40%", "h": "30%",
             "steps": ["按Temp设温度(一般4°C)", "按Speed设转速(看转子最大限制)", "按Time设时间", "按Start启动"],
             "tips": ["必须配平！对称管重量差<0.01g", "必须等转速归零再开盖", "超过转子最大转速会炸裂"],
             "video": {"search": "高速冷冻离心机 操作教程"}},
            {"name": "转子仓", "desc": "安装转子和样品管", "x": "10%", "y": "50%", "w": "80%", "h": "40%",
             "steps": ["检查转子安装牢固", "对称放置样品管", "用天平配平(差<0.01g)", "关盖锁紧"],
             "tips": ["不配平→剧烈震动→仪器损坏", "管不能装太满(最多2/3)", "转子有最大转速限制，看转子上的标注"],
             "video": {"search": "离心机 转子 安装 配平"}},
        ],
    },
    {
        "id": "I002", "name": "紫外分光光度计",
        "icon": "spectrophotometer",
        "desc": "测溶液吸光度，算出目标物质浓度",
        "protocols": ["P001", "P002", "P003", "P004", "P005", "P011", "P012", "P013", "P014", "P015", "P016", "P017"],
        "color": "#2E7D32",
        "hotspots": [
            {"name": "样品仓", "desc": "放入比色皿", "x": "35%", "y": "30%", "w": "30%", "h": "25%",
             "steps": ["打开样品仓盖", "放入比色皿(透光面朝光源方向)", "关盖"],
             "tips": ["比色皿透光面不能用手摸！", "用擦镜纸擦拭", "拿比色皿只捏磨砂面"],
             "video": {"search": "分光光度计 比色皿 使用方法"}},
            {"name": "操作面板", "desc": "设置波长、调零、读数", "x": "10%", "y": "10%", "w": "80%", "h": "20%",
             "steps": ["开机预热20min", "设置目标波长(如620nm)", "空白对照调零(Run/Zero)", "放入样品读数"],
             "tips": ["必须预热20min以上！", "空白对照每次换波长都要重调", "OD值超出0.1-0.8需稀释样品"],
             "video": {"search": "紫外分光光度计 操作教程"}},
        ],
    },
    {
        "id": "I003", "name": "PCR仪",
        "icon": "pcr_machine",
        "desc": "给DNA做复印机，通过温度循环扩增目标片段",
        "protocols": ["P007"],
        "color": "#00838F",
        "hotspots": [
            {"name": "热盖", "desc": "加热盖子防蒸发", "x": "20%", "y": "10%", "w": "60%", "h": "20%",
             "steps": ["确保热盖温度设为105°C", "PCR管盖要压紧"],
             "tips": ["热盖温度不够→管盖内壁冷凝→反应体积变化"],
             "video": {"search": "PCR仪 操作教程 设置程序"}},
            {"name": "样品槽", "desc": "放置PCR管", "x": "25%", "y": "40%", "w": "50%", "h": "40%",
             "steps": ["将PCR管放入孔中", "确保管底和金属孔充分接触", "空孔用空管填充"],
             "tips": ["管底有气泡→传热不均→扩增不均", "不要用半裙边板(高度不匹配)"],
             "video": {"search": "PCR仪 上机 操作步骤"}},
        ],
    },
    {
        "id": "I004", "name": "高压灭菌锅",
        "icon": "autoclave",
        "desc": "121°C高压杀灭所有微生物",
        "protocols": ["P006", "P007", "P010", "P018"],
        "color": "#D84315",
        "hotspots": [
            {"name": "锅盖", "desc": "密封加压", "x": "20%", "y": "5%", "w": "60%", "h": "25%",
             "steps": ["对角线拧紧螺丝", "检查密封圈完好"],
             "tips": ["密封圈老化→漏气→温度达不到", "螺丝没拧紧→蒸汽喷出→烫伤"],
             "video": {"search": "高压灭菌锅 操作教程 使用方法"}},
            {"name": "控制面板", "desc": "设置灭菌程序", "x": "60%", "y": "30%", "w": "30%", "h": "40%",
             "steps": ["设温度121°C", "设时间15-20min(液体30min)", "按Start"],
             "tips": ["液体不能装太满(最多2/3)", "液体灭菌完要自然降压，不能快排"],
             "video": {"search": "高压灭菌锅 灭菌程序 设置"}},
        ],
    },
    {
        "id": "I005", "name": "pH计",
        "icon": "ph_meter",
        "desc": "测溶液酸碱度",
        "protocols": ["P001", "P004", "P005", "P006", "P010", "P011", "P012", "P015"],
        "color": "#6A1B9A",
        "hotspots": [
            {"name": "电极", "desc": "pH感应探头", "x": "40%", "y": "10%", "w": "20%", "h": "60%",
             "steps": ["电极必须泡在3mol/L KCl保护液中保存", "用前用蒸馏水冲洗", "用滤纸轻轻吸干(不能擦！)"],
             "tips": ["电极干燥→读数不准且损坏电极", "不能擦电极球泡→产生静电→读数漂移"],
             "video": {"search": "pH计 电极 保存 使用"}},
            {"name": "操作面板", "desc": "校准和读数", "x": "10%", "y": "15%", "w": "30%", "h": "50%",
             "steps": ["两点校准：先pH6.86，再pH4.00(或9.18)", "校准后测样品"],
             "tips": ["每次用前必须校准！", "校准液要新鲜(3个月换一次)"],
             "video": {"search": "pH计 校准 两点校准 操作"}},
        ],
    },
    {
        "id": "I006", "name": "恒温水浴锅",
        "icon": "water_bath",
        "desc": "给反应提供恒定温度环境",
        "protocols": ["P001", "P003", "P006", "P013", "P014", "P015"],
        "color": "#0277BD",
        "hotspots": [
            {"name": "控制面板", "desc": "设置温度", "x": "10%", "y": "10%", "w": "35%", "h": "30%",
             "steps": ["设置目标温度", "等待温度稳定(约10min)", "用水银温度计校验"],
             "tips": ["面板显示温度可能不准→用水银温度计校验", "水位不能低于加热管"],
             "video": {"search": "恒温水浴锅 使用方法 操作"}},
        ],
    },
    {
        "id": "I007", "name": "凝胶成像系统",
        "icon": "gel_doc",
        "desc": "给跑完电泳的凝胶拍照，看DNA条带",
        "protocols": ["P008"],
        "color": "#37474F",
        "hotspots": [
            {"name": "紫外透射台", "desc": "放置凝胶并紫外激发", "x": "20%", "y": "30%", "w": "60%", "h": "50%",
             "steps": ["戴护目镜！", "将凝胶放在透射台上", "关上暗箱门", "打开紫外灯"],
             "tips": ["紫外灯开启时不能直视！", "凝胶放反了→条带左右颠倒", "拍照后立即关紫外灯"],
             "video": {"search": "凝胶成像系统 操作 拍照"}},
        ],
    },
    {
        "id": "I008", "name": "电泳槽",
        "icon": "electrophoresis",
        "desc": "让DNA在电场中按大小分开",
        "protocols": ["P008"],
        "color": "#4E342E",
        "hotspots": [
            {"name": "凝胶槽", "desc": "放置凝胶和缓冲液", "x": "15%", "y": "20%", "w": "70%", "h": "60%",
             "steps": ["将凝胶放入槽中(靠近黑色负极端)", "加缓冲液没过凝胶2mm", "拔梳子(垂直向上)", "接电源(红正黑负)"],
             "tips": ["正负极接反→DNA往反方向跑！", "缓冲液不够→局部干燥→条带扭曲", "拔梳子歪了→孔变形→条带歪"],
             "video": {"search": "琼脂糖凝胶电泳 操作教程"}},
        ],
    },
    {
        "id": "I009", "name": "液氮罐",
        "icon": "ln2_tank",
        "desc": "存放-196°C液氮，冻存样品",
        "protocols": ["P001", "P002", "P003", "P004", "P005", "P006", "P009", "P010", "P011", "P012", "P013", "P014", "P015", "P016", "P017", "P018"],
        "color": "#0D47A1",
        "hotspots": [
            {"name": "罐口", "desc": "取放样品", "x": "30%", "y": "5%", "w": "40%", "h": "20%",
             "steps": ["戴防冻手套+护目镜", "用长勺取放样品", "操作要快减少液氮蒸发"],
             "tips": ["裸手碰液氮→严重冻伤！", "密闭空间液氮蒸发→缺氧窒息", "液氮溅入眼睛→永久损伤"],
             "video": {"search": "液氮罐 使用 取样 操作"}},
        ],
    },
    {
        "id": "I010", "name": "电子天平",
        "icon": "balance",
        "desc": "精确称量试剂",
        "protocols": ["P001", "P002", "P010"],
        "color": "#558B2F",
        "hotspots": [
            {"name": "称量盘", "desc": "放置称量纸和试剂", "x": "25%", "y": "25%", "w": "50%", "h": "40%",
             "steps": ["水平气泡必须在中间", "放称量纸或称量瓶", "按Tare去皮", "用药匙加试剂"],
             "tips": ["气泡不在中间→称量不准→调地脚螺丝", "不能直接把试剂倒在称量盘上", "有腐蚀性试剂必须用称量瓶"],
             "video": {"search": "电子天平 使用方法 称量"}},
            {"name": "操作面板", "desc": "读数和功能键", "x": "10%", "y": "10%", "w": "30%", "h": "30%",
             "steps": ["开机预热", "按Cal校准(放标准砝码)", "按Tare去皮"],
             "tips": ["每天用前校准！", "精度0.0001g的天平要防震"],
             "video": {"search": "电子天平 校准 去皮 操作"}},
        ],
    },
]

# ========== 试剂分子量数据库 ==========
MW_DB = {
    "葡萄糖": 180.16, "蔗糖": 342.30, "Tris": 121.14, "EDTA": 372.24,
    "EDTA-2Na": 336.21, "NaCl": 58.44, "KCl": 74.55, "MgCl2": 95.21,
    "CaCl2": 110.98, "NaOH": 40.00, "KOH": 56.11, "HCl": 36.46,
    "H2SO4": 98.08, "SDS": 288.38, "DTT": 154.25, "VC(抗坏血酸)": 176.12,
    "愈创木酚": 124.14, "TBA": 144.15, "TCA": 163.39, "蒽酮": 194.23,
    "BSA": 66430, "NaH2PO4": 119.98, "Na2HPO4": 141.96, "KH2PO4": 136.09,
    "柠檬酸": 192.12, "柠檬酸钠": 258.06, "CTAB": 364.45, "甘油": 92.09,
}


# ========== Protocol文档解析 ==========
def parse_protocol(filepath):
    """解析Protocol txt文件，按section返回结构化内容"""
    if not os.path.exists(filepath):
        return None
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()

    sections = {
        "meta": {}, "principle": "", "source": "", "instruments": "",
        "reagents": "", "formula": "", "steps": [], "safety": "",
        "tips": "", "data": "", "troubleshoot": "",
    }

    current_section = None
    lines = text.split("\n")
    buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("Protocol编号"):
            sections["meta"]["id"] = stripped.split("：")[-1].strip() if "：" in stripped else stripped.split(":")[-1].strip()
        elif stripped.startswith("实验名称"):
            sections["meta"]["name"] = stripped.split("：")[-1].strip() if "：" in stripped else stripped.split(":")[-1].strip()
        elif stripped.startswith("一、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "principle"
            buffer = []
        elif stripped.startswith("二、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "principle_inner"
            buffer = []
        elif stripped.startswith("三、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "source"
            buffer = []
        elif stripped.startswith("四、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "instruments"
            buffer = []
        elif stripped.startswith("五、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "reagents"
            buffer = []
        elif stripped.startswith("六、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "formula"
            buffer = []
        elif stripped.startswith("七、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "steps"
            buffer = []
        elif stripped.startswith("八、"):
            if current_section and buffer:
                if current_section == "steps":
                    sections["steps_raw"] = "\n".join(buffer)
                else:
                    sections[current_section] = "\n".join(buffer)
            current_section = "safety"
            buffer = []
        elif stripped.startswith("九、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "tips"
            buffer = []
        elif stripped.startswith("十、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "data"
            buffer = []
        elif stripped.startswith("十一、"):
            if current_section and buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = "troubleshoot"
            buffer = []
        else:
            buffer.append(line)

    if current_section and buffer:
        if current_section == "steps":
            sections["steps_raw"] = "\n".join(buffer)
        else:
            sections[current_section] = "\n".join(buffer)

    # Parse steps into structured list
    steps_raw = sections.get("steps_raw", sections.get("steps", ""))
    if isinstance(steps_raw, str):
        parsed_steps = []
        current_step = None
        for line in steps_raw.split("\n"):
            stripped = line.strip()
            if stripped.startswith("步骤"):
                if current_step:
                    parsed_steps.append(current_step)
                current_step = {"title": stripped, "why": "", "how": "", "correct": "", "wrong": ""}
            elif current_step:
                if stripped.startswith("为什么"):
                    current_step["why"] = stripped.split("：")[-1].strip() if "：" in stripped else stripped[3:]
                elif stripped.startswith("做对了"):
                    current_step["correct"] = stripped.split("：")[-1].strip() if "：" in stripped else stripped[3:]
                elif stripped.startswith("做错了"):
                    current_step["wrong"] = stripped.split("：")[-1].strip() if "：" in stripped else stripped[3:]
                elif not any(stripped.startswith(k) for k in ["为什么", "做对了", "做错了", "步骤"]):
                    if current_step["how"]:
                        current_step["how"] += " " + stripped
                    else:
                        current_step["how"] = stripped
        if current_step:
            parsed_steps.append(current_step)
        sections["steps"] = parsed_steps

    return sections


# ========== 仪器SVG图形生成 ==========
INSTRUMENT_SVG = {
    "centrifuge": """<svg viewBox="0 0 120 120"><rect x="25" y="20" width="70" height="80" rx="10" fill="#E3F2FD"/><rect x="30" y="25" width="60" height="35" rx="8" fill="#1565C0"/><circle cx="60" cy="42" r="14" fill="white" opacity="0.3"/><circle cx="60" cy="42" r="10" fill="white" opacity="0.5"/><circle cx="60" cy="42" r="4" fill="#0D47A1"/><line x1="60" y1="42" x2="50" y2="35" stroke="#0D47A1" stroke-width="2"/><line x1="60" y1="42" x2="70" y2="35" stroke="#0D47A1" stroke-width="2"/><line x1="60" y1="42" x2="60" y2="32" stroke="#0D47A1" stroke-width="2"/><rect x="35" y="70" width="50" height="20" rx="4" fill="#90CAF9"/><rect x="42" y="74" width="8" height="12" rx="2" fill="#1565C0"/><circle cx="58" cy="80" r="3" fill="#1565C0"/><circle cx="68" cy="80" r="3" fill="#1565C0"/></svg>""",
    "spectrophotometer": """<svg viewBox="0 0 120 120"><rect x="15" y="30" width="90" height="70" rx="8" fill="#E8F5E9"/><rect x="20" y="35" width="40" height="30" rx="4" fill="#2E7D32"/><rect x="25" y="40" width="30" height="20" rx="2" fill="#C8E6C9"/><rect x="70" y="38" width="30" height="24" rx="3" fill="#F5F5F5" stroke="#BDBDBD"/><text x="85" y="54" text-anchor="middle" font-size="8" fill="#333" font-weight="bold">0.452</text><rect x="55" y="70" width="15" height="25" rx="3" fill="#FFF9C4" stroke="#F9A825"/><rect x="25" y="80" width="8" height="5" rx="1" fill="#FF6F00"/><rect x="38" y="80" width="8" height="5" rx="1" fill="#FF6F00"/><rect x="51" y="80" width="8" height="5" rx="1" fill="#FF6F00"/><rect x="20" y="95" width="80" height="5" rx="2" fill="#E0E0E0"/></svg>""",
    "pcr_machine": """<svg viewBox="0 0 120 120"><rect x="20" y="25" width="80" height="70" rx="10" fill="#E0F2F1"/><rect x="25" y="30" width="70" height="20" rx="6" fill="#00838F"/><rect x="30" y="55" width="60" height="30" rx="4" fill="#B2DFDB"/><rect x="35" y="60" width="12" height="8" rx="2" fill="#00838F"/><rect x="52" y="60" width="12" height="8" rx="2" fill="#00838F"/><rect x="69" y="60" width="12" height="8" rx="2" fill="#00838F"/><rect x="35" y="72" width="12" height="8" rx="2" fill="#00838F"/><rect x="52" y="72" width="12" height="8" rx="2" fill="#00838F"/><rect x="69" y="72" width="12" height="8" rx="2" fill="#00838F"/></svg>""",
    "autoclave": """<svg viewBox="0 0 120 120"><ellipse cx="60" cy="85" rx="35" ry="15" fill="#B0BEC5"/><rect x="25" y="35" width="70" height="55" rx="8" fill="#CFD8DC"/><ellipse cx="60" cy="35" rx="35" ry="12" fill="#ECEFF1"/><circle cx="60" cy="35" r="8" fill="#FF5722" opacity="0.3"/><circle cx="60" cy="35" r="4" fill="#FF5722"/><rect x="42" y="25" width="6" height="12" rx="2" fill="#78909C"/><rect x="57" y="25" width="6" height="12" rx="2" fill="#78909C"/><rect x="72" y="25" width="6" height="12" rx="2" fill="#78909C"/><circle cx="40" cy="65" r="5" fill="#FFF" stroke="#FF5722"/><text x="40" y="68" text-anchor="middle" font-size="6" fill="#FF5722" font-weight="bold">121°</text><circle cx="80" cy="65" r="5" fill="#FFF" stroke="#4CAF50"/><text x="80" y="68" text-anchor="middle" font-size="6" fill="#4CAF50" font-weight="bold">OK</text></svg>""",
    "ph_meter": """<svg viewBox="0 0 120 120"><rect x="35" y="20" width="50" height="80" rx="8" fill="#EDE7F6"/><rect x="40" y="25" width="40" height="30" rx="4" fill="#7B1FA2"/><rect x="45" y="30" width="30" height="20" rx="2" fill="#E1BEE7"/><text x="60" y="44" text-anchor="middle" font-size="10" fill="white" font-weight="bold">7.00</text><rect x="42" y="62" width="36" height="8" rx="2" fill="#CE93D8"/><rect x="55" y="75" width="10" height="30" rx="3" fill="#9E9E9E"/><circle cx="60" cy="105" r="5" fill="#7B1FA2"/></svg>""",
    "water_bath": """<svg viewBox="0 0 120 120"><rect x="15" y="40" width="90" height="60" rx="8" fill="#E3F2FD"/><rect x="20" y="45" width="80" height="40" rx="6" fill="#BBDEFB"/><ellipse cx="60" cy="55" rx="30" ry="6" fill="#90CAF9" opacity="0.5"/><circle cx="50" cy="65" r="3" fill="#42A5F5" opacity="0.4"/><circle cx="70" cy="60" r="2" fill="#42A5F5" opacity="0.4"/><circle cx="60" cy="70" r="2.5" fill="#42A5F5" opacity="0.4"/><rect x="20" y="90" width="25" height="10" rx="3" fill="#1565C0"/><text x="32" y="98" text-anchor="middle" font-size="7" fill="white">100°C</text><rect x="55" y="90" width="45" height="10" rx="3" fill="#E0E0E0"/></svg>""",
    "gel_doc": """<svg viewBox="0 0 120 120"><rect x="20" y="25" width="80" height="70" rx="10" fill="#37474F"/><rect x="28" y="33" width="64" height="45" rx="6" fill="#263238"/><rect x="35" y="40" width="50" height="30" rx="3" fill="#1B5E20" opacity="0.8"/><rect x="42" y="45" width="3" height="20" fill="#76FF03" opacity="0.7"/><rect x="50" y="48" width="3" height="14" fill="#76FF03" opacity="0.6"/><rect x="58" y="44" width="3" height="22" fill="#76FF03" opacity="0.8"/><rect x="66" y="50" width="3" height="12" fill="#76FF03" opacity="0.5"/><rect x="74" y="46" width="3" height="18" fill="#76FF03" opacity="0.7"/><rect x="35" y="85" width="50" height="8" rx="2" fill="#546E7A"/><circle cx="42" cy="89" r="2" fill="#4CAF50"/><circle cx="52" cy="89" r="2" fill="#FF9800"/><circle cx="62" cy="89" r="2" fill="#F44336"/></svg>""",
    "electrophoresis": """<svg viewBox="0 0 120 120"><rect x="15" y="35" width="90" height="65" rx="6" fill="#EFEBE9"/><rect x="20" y="40" width="80" height="50" rx="4" fill="#D7CCC8"/><rect x="30" y="50" width="60" height="15" rx="2" fill="#FFF9C4"/><rect x="35" y="45" width="4" height="25" fill="#795548"/><rect x="75" y="45" width="4" height="25" fill="#795548"/><line x1="20" y1="48" x2="35" y2="48" stroke="#F44336" stroke-width="2"/><line x1="85" y1="48" x2="100" y2="48" stroke="#212121" stroke-width="2"/><rect x="35" y="55" width="3" height="8" fill="#212121" opacity="0.3"/><rect x="45" y="53" width="3" height="12" fill="#212121" opacity="0.5"/><rect x="55" y="56" width="3" height="6" fill="#212121" opacity="0.2"/><rect x="65" y="54" width="3" height="10" fill="#212121" opacity="0.4"/><text x="40" y="75" font-size="6" fill="#795548" font-weight="bold">+ (red)</text><text x="75" y="75" font-size="6" fill="#795548" font-weight="bold">- (black)</text></svg>""",
    "ln2_tank": """<svg viewBox="0 0 120 120"><ellipse cx="60" cy="95" rx="28" ry="10" fill="#90CAF9"/><rect x="32" y="25" width="56" height="75" rx="6" fill="#E3F2FD"/><rect x="36" y="30" width="48" height="65" rx="4" fill="#BBDEFB"/><rect x="45" y="18" width="30" height="15" rx="6" fill="#1565C0"/><rect x="55" y="15" width="10" height="8" rx="3" fill="#90CAF9"/><circle cx="50" cy="60" r="6" fill="#90CAF9" opacity="0.5"/><circle cx="65" cy="55" r="5" fill="#90CAF9" opacity="0.4"/><circle cx="58" cy="70" r="4" fill="#90CAF9" opacity="0.3"/><text x="60" y="88" text-anchor="middle" font-size="7" fill="#0D47A1" font-weight="bold">-196°C</text></svg>""",
    "balance": """<svg viewBox="0 0 120 120"><rect x="25" y="65" width="70" height="35" rx="6" fill="#E8EAF6"/><rect x="30" y="70" width="60" height="25" rx="4" fill="#C5CAE9"/><rect x="35" y="72" width="40" height="15" rx="2" fill="#FFF" stroke="#E0E0E0"/><text x="55" y="83" text-anchor="middle" font-size="9" fill="#333" font-weight="bold">0.0000</text><rect x="80" y="75" width="8" height="4" rx="1" fill="#FF6F00"/><rect x="80" y="82" width="8" height="4" rx="1" fill="#4CAF50"/><rect x="35" y="55" width="50" height="12" rx="3" fill="#F5F5F5" stroke="#BDBDBD"/><rect x="38" y="58" width="44" height="6" rx="2" fill="#FFF"/><line x1="60" y1="60" x2="60" y2="65" stroke="#9E9E9E" stroke-width="1"/></svg>""",
}

def get_instrument_svg(icon_type):
    """获取仪器SVG图形"""
    return INSTRUMENT_SVG.get(icon_type, "")


# ========== Protocol内容缓存 ==========
_PROTOCOL_CONTENT_CACHE = {}

def get_protocol_content(protocol_id):
    """获取Protocol文本内容（带缓存）"""
    if protocol_id not in _PROTOCOL_CONTENT_CACHE:
        meta = next((p for p in PROTOCOL_META if p["id"] == protocol_id), None)
        if not meta:
            return ""
        filepath = os.path.join(PROTOCOL_DIR, meta["file"])
        if not os.path.exists(filepath):
            return ""
        with open(filepath, "r", encoding="utf-8") as f:
            _PROTOCOL_CONTENT_CACHE[protocol_id] = f.read()
    return _PROTOCOL_CONTENT_CACHE[protocol_id]


# ========== 认证路由 ==========

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        conn = get_user_db()
        user = conn.execute("SELECT * FROM users WHERE username=? OR email=?",
                            (username, username)).fetchone()
        if user and check_password_hash(user["password_hash"], password):
            session["user_id"] = user["id"]
            session["username"] = user["username"]
            session["role"] = user["role"]
            conn.execute("UPDATE users SET last_login=? WHERE id=?", (datetime.now(), user["id"]))
            conn.commit()
            conn.close()
            flash(f"欢迎回来，{user['username']}！", "success")
            return redirect(url_for("home"))
        conn.close()
        flash("用户名或密码错误", "error")
    return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm = request.form.get("confirm", "")

        if len(username) < 2:
            flash("用户名至少2个字符", "error")
        elif len(password) < 6:
            flash("密码至少6位", "error")
        elif password != confirm:
            flash("两次密码不一致", "error")
        else:
            conn = get_user_db()
            existing = conn.execute("SELECT id FROM users WHERE username=? OR email=?",
                                    (username, email)).fetchone()
            if existing:
                flash("用户名或邮箱已被注册", "error")
            else:
                conn.execute(
                    "INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)",
                    (username, email, generate_password_hash(password))
                )
                conn.commit()
                conn.close()
                flash("注册成功，请登录", "success")
                return redirect(url_for("login"))
            conn.close()
    return render_template("register.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("已退出登录", "info")
    return redirect(url_for("home"))


# ========== 小程序API ==========
@app.route("/api/wx-login", methods=["POST"])
def wx_login():
    """小程序登录 - 用户名密码方式"""
    data = request.json or {}
    username = data.get("username", "").strip()
    password = data.get("password", "")
    if not username or not password:
        return jsonify({"error": "请输入用户名和密码"})
    conn = get_user_db()
    user = conn.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()
    conn.close()
    if not user or not check_password_hash(user["password_hash"], password):
        return jsonify({"error": "用户名或密码错误"})
    token = _generate_token(user["id"])
    return jsonify({
        "token": token,
        "user": {"id": user["id"], "username": user["username"], "role": user["role"]}
    })


@app.route("/api/wx-register", methods=["POST"])
def wx_register():
    """小程序注册"""
    data = request.json or {}
    username = data.get("username", "").strip()
    email = data.get("email", "").strip()
    password = data.get("password", "")
    if not username or not email or not password:
        return jsonify({"error": "请填写完整信息"})
    conn = get_user_db()
    existing = conn.execute("SELECT id FROM users WHERE username=? OR email=?", (username, email)).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": "用户名或邮箱已存在"})
    conn.execute("INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)",
                 (username, email, generate_password_hash(password)))
    conn.commit()
    user = conn.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()
    conn.close()
    token = _generate_token(user["id"])
    return jsonify({
        "token": token,
        "user": {"id": user["id"], "username": user["username"], "role": user["role"]}
    })


@app.route("/api/wx-data")
@login_required
def wx_data():
    """获取用户所有数据"""
    uid = session["user_id"]
    conn = get_user_db()
    # 收藏
    favs = [r["protocol_id"] for r in conn.execute(
        "SELECT protocol_id FROM protocol_favorites WHERE user_id=?", (uid,)).fetchall()]
    # 评分
    ratings = {}
    for r in conn.execute("SELECT protocol_id, rating, comment FROM protocol_ratings WHERE user_id=?", (uid,)).fetchall():
        ratings[r["protocol_id"]] = {"rating": r["rating"], "comment": r["comment"]}
    # 日志
    logs = []
    for r in conn.execute("SELECT * FROM experiment_logs WHERE user_id=? ORDER BY id DESC", (uid,)).fetchall():
        logs.append({"id": r["id"], "title": r["title"], "protocol_id": r["protocol_id"],
                      "content": r["content"], "tags": r["tags"], "status": r["status"],
                      "created_at": r["created_at"]})
    conn.close()
    return jsonify({"favorites": favs, "ratings": ratings, "logs": logs})


@app.route("/api/wx-sync", methods=["POST"])
@login_required
def wx_sync():
    """批量同步小程序数据到服务端"""
    uid = session["user_id"]
    data = request.json or {}
    conn = get_user_db()
    synced = {"favorites": 0, "ratings": 0, "logs": 0}

    # 同步收藏
    for pid in data.get("favorites", []):
        try:
            conn.execute("INSERT OR IGNORE INTO protocol_favorites (user_id, protocol_id) VALUES (?, ?)", (uid, pid))
            synced["favorites"] += 1
        except: pass

    # 同步评分
    for pid, info in data.get("ratings", {}).items():
        try:
            conn.execute(
                "INSERT INTO protocol_ratings (user_id, protocol_id, rating, comment) VALUES (?, ?, ?, ?) "
                "ON CONFLICT(user_id, protocol_id) DO UPDATE SET rating=excluded.rating, comment=excluded.comment",
                (uid, pid, info.get("rating", 5), info.get("comment", "")))
            synced["ratings"] += 1
        except: pass

    # 同步日志
    for log in data.get("logs", []):
        try:
            conn.execute(
                "INSERT INTO experiment_logs (user_id, title, protocol_id, content, tags, status) VALUES (?, ?, ?, ?, ?, ?)",
                (uid, log.get("title", ""), log.get("protocolId", ""), log.get("content", ""),
                 log.get("tags", ""), log.get("status", "done")))
            synced["logs"] += 1
        except: pass

    conn.commit()
    conn.close()
    return jsonify({"synced": synced})


@app.route("/profile")
@login_required
def profile():
    conn = get_user_db()
    user = conn.execute("SELECT * FROM users WHERE id=?", (session["user_id"],)).fetchone()
    history = conn.execute(
        "SELECT * FROM user_data WHERE user_id=? ORDER BY created_at DESC LIMIT 20",
        (session["user_id"],)
    ).fetchall()
    api_row = conn.execute("SELECT * FROM user_api_settings WHERE user_id=?", (session["user_id"],)).fetchone()
    conn.close()
    api_settings = {
        "provider": api_row["provider"] if api_row else "anthropic",
        "api_key": api_row["api_key"] if api_row else "",
        "model": api_row["model"] if api_row else "claude-sonnet-4-20250514",
        "base_url": api_row["base_url"] if api_row else "",
    } if api_row else None
    return render_template("profile.html", user=user, history=history, api_settings=api_settings)


@app.route("/api/user/save-data", methods=["POST"])
@login_required
def save_user_data():
    data = request.json
    conn = get_user_db()
    conn.execute(
        "INSERT INTO user_data (user_id, data_type, data_name, data_json) VALUES (?, ?, ?, ?)",
        (session["user_id"], data["type"], data["name"], json.dumps(data["content"], ensure_ascii=False))
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/user/history")
@login_required
def user_history():
    conn = get_user_db()
    rows = conn.execute(
        "SELECT * FROM user_data WHERE user_id=? ORDER BY created_at DESC LIMIT 50",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/user/delete-data/<int:data_id>", methods=["POST"])
@login_required
def delete_user_data(data_id):
    conn = get_user_db()
    conn.execute("DELETE FROM user_data WHERE id=? AND user_id=?", (data_id, session["user_id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ========== 用户API设置 ==========

@app.route("/api/user/api-settings", methods=["GET"])
@login_required
def get_api_settings():
    conn = get_user_db()
    row = conn.execute("SELECT * FROM user_api_settings WHERE user_id=?", (session["user_id"],)).fetchone()
    conn.close()
    if row:
        return jsonify({
            "provider": row["provider"],
            "api_key": row["api_key"],
            "model": row["model"],
            "base_url": row["base_url"] or "",
        })
    return jsonify({"provider": "anthropic", "api_key": "", "model": "claude-sonnet-4-20250514", "base_url": ""})


@app.route("/api/user/api-settings", methods=["POST"])
@login_required
def save_api_settings():
    data = request.json
    provider = data.get("provider", "anthropic")
    api_key = data.get("api_key", "")
    model = data.get("model", "")
    base_url = data.get("base_url", "")

    # 根据provider设置默认模型
    default_models = {
        "anthropic": "claude-sonnet-4-20250514",
        "openai": "gpt-4o",
        "deepseek": "deepseek-v4-pro",
        "custom": "",
    }
    if not model:
        model = default_models.get(provider, "")

    conn = get_user_db()
    existing = conn.execute("SELECT id FROM user_api_settings WHERE user_id=?", (session["user_id"],)).fetchone()
    if existing:
        conn.execute(
            "UPDATE user_api_settings SET provider=?, api_key=?, model=?, base_url=?, updated_at=? WHERE user_id=?",
            (provider, api_key, model, base_url, datetime.now(), session["user_id"])
        )
    else:
        conn.execute(
            "INSERT INTO user_api_settings (user_id, provider, api_key, model, base_url) VALUES (?, ?, ?, ?, ?)",
            (session["user_id"], provider, api_key, model, base_url)
        )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/user/update-profile", methods=["POST"])
@login_required
def update_profile():
    data = request.json
    username = data.get("username", "").strip()
    email = data.get("email", "").strip()

    if len(username) < 2:
        return jsonify({"error": "用户名至少2个字符"}), 400

    conn = get_user_db()
    # 检查用户名是否被其他人占用
    existing = conn.execute(
        "SELECT id FROM users WHERE username=? AND id!=?",
        (username, session["user_id"])
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": "用户名已被占用"}), 400

    # 检查邮箱是否被其他人占用
    existing_email = conn.execute(
        "SELECT id FROM users WHERE email=? AND id!=?",
        (email, session["user_id"])
    ).fetchone()
    if existing_email:
        conn.close()
        return jsonify({"error": "邮箱已被占用"}), 400

    conn.execute("UPDATE users SET username=?, email=? WHERE id=?",
                 (username, email, session["user_id"]))
    session["username"] = username
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/user/change-password", methods=["POST"])
@login_required
def change_password():
    data = request.json
    old_pw = data.get("old_password", "")
    new_pw = data.get("new_password", "")

    if len(new_pw) < 6:
        return jsonify({"error": "新密码至少6位"}), 400

    conn = get_user_db()
    user = conn.execute("SELECT password_hash FROM users WHERE id=?", (session["user_id"],)).fetchone()
    if not user or not check_password_hash(user["password_hash"], old_pw):
        conn.close()
        return jsonify({"error": "原密码错误"}), 400

    conn.execute("UPDATE users SET password_hash=? WHERE id=?",
                 (generate_password_hash(new_pw), session["user_id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ========== Protocol收藏/评分API ==========

@app.route("/api/favorite/<pid>", methods=["POST"])
@login_required
def toggle_favorite(pid):
    conn = get_user_db()
    existing = conn.execute(
        "SELECT id FROM protocol_favorites WHERE user_id=? AND protocol_id=?",
        (session["user_id"], pid)
    ).fetchone()
    if existing:
        conn.execute("DELETE FROM protocol_favorites WHERE id=?", (existing["id"],))
        conn.commit()
        conn.close()
        return jsonify({"favorited": False})
    else:
        conn.execute(
            "INSERT INTO protocol_favorites (user_id, protocol_id) VALUES (?, ?)",
            (session["user_id"], pid)
        )
        conn.commit()
        conn.close()
        return jsonify({"favorited": True})


@app.route("/api/favorites")
@login_required
def get_favorites():
    conn = get_user_db()
    rows = conn.execute(
        "SELECT protocol_id FROM protocol_favorites WHERE user_id=?",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    return jsonify([r["protocol_id"] for r in rows])


@app.route("/api/rate/<pid>", methods=["POST"])
@login_required
def rate_protocol(pid):
    data = request.json
    rating = int(data.get("rating", 5))
    comment = data.get("comment", "")
    conn = get_user_db()
    existing = conn.execute(
        "SELECT id FROM protocol_ratings WHERE user_id=? AND protocol_id=?",
        (session["user_id"], pid)
    ).fetchone()
    if existing:
        conn.execute(
            "UPDATE protocol_ratings SET rating=?, comment=? WHERE id=?",
            (rating, comment, existing["id"])
        )
    else:
        conn.execute(
            "INSERT INTO protocol_ratings (user_id, protocol_id, rating, comment) VALUES (?, ?, ?, ?)",
            (session["user_id"], pid, rating, comment)
        )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/ratings/<pid>")
def get_ratings(pid):
    conn = get_user_db()
    avg = conn.execute(
        "SELECT AVG(rating) as avg_rating, COUNT(*) as count FROM protocol_ratings WHERE protocol_id=?",
        (pid,)
    ).fetchone()
    my_rating = None
    if "user_id" in session:
        r = conn.execute(
            "SELECT rating FROM protocol_ratings WHERE user_id=? AND protocol_id=?",
            (session["user_id"], pid)
        ).fetchone()
        if r:
            my_rating = r["rating"]
    recent = conn.execute(
        "SELECT r.rating, r.comment, r.created_at, u.username FROM protocol_ratings r "
        "JOIN users u ON r.user_id=u.id WHERE r.protocol_id=? ORDER BY r.created_at DESC LIMIT 10",
        (pid,)
    ).fetchall()
    conn.close()
    return jsonify({
        "avg": round(avg["avg_rating"], 1) if avg["avg_rating"] else 0,
        "count": avg["count"],
        "my_rating": my_rating,
        "recent": [dict(r) for r in recent]
    })


@app.route("/api/top-protocols")
def top_protocols():
    conn = get_user_db()
    favs = conn.execute(
        "SELECT protocol_id, COUNT(*) as cnt FROM protocol_favorites GROUP BY protocol_id ORDER BY cnt DESC LIMIT 6"
    ).fetchall()
    rated = conn.execute(
        "SELECT protocol_id, AVG(rating) as avg_r, COUNT(*) as cnt FROM protocol_ratings "
        "GROUP BY protocol_id HAVING cnt>=2 ORDER BY avg_r DESC LIMIT 6"
    ).fetchall()
    conn.close()
    return jsonify({"favorites": [dict(r) for r in favs], "rated": [dict(r) for r in rated]})


# ========== 实验日志API ==========

@app.route("/journal")
@login_required
def journal_page():
    return render_template("journal.html")


@app.route("/api/journal", methods=["GET"])
@login_required
def get_journal():
    conn = get_user_db()
    logs = conn.execute(
        "SELECT * FROM experiment_logs WHERE user_id=? ORDER BY created_at DESC LIMIT 100",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in logs])


@app.route("/api/journal", methods=["POST"])
@login_required
def add_journal():
    data = request.json
    conn = get_user_db()
    conn.execute(
        "INSERT INTO experiment_logs (user_id, title, protocol_id, content, tags, status) VALUES (?, ?, ?, ?, ?, ?)",
        (session["user_id"], data["title"], data.get("protocol_id", ""),
         data["content"], data.get("tags", ""), data.get("status", "done"))
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/journal/<int:lid>", methods=["DELETE"])
@login_required
def delete_journal(lid):
    conn = get_user_db()
    conn.execute("DELETE FROM experiment_logs WHERE id=? AND user_id=?", (lid, session["user_id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/pricing")
def pricing():
    return render_template("pricing.html")


@app.route("/teams")
def teams_page():
    return render_template("teams.html")


# ========== 团队协作API ==========

@app.route("/api/teams", methods=["GET"])
@login_required
def api_teams_list():
    """获取用户的所有团队"""
    uid = session["user_id"]
    conn = get_user_db()
    teams = conn.execute("""
        SELECT t.*, (SELECT COUNT(*) FROM team_members WHERE team_id=t.id) as member_count
        FROM teams t
        WHERE t.owner_id=? OR t.id IN (SELECT team_id FROM team_members WHERE user_id=?)
        ORDER BY t.created_at DESC
    """, (uid, uid)).fetchall()
    conn.close()
    return jsonify({"teams": [dict(t) for t in teams]})


@app.route("/api/teams", methods=["POST"])
@login_required
def api_teams_create():
    """创建团队"""
    data = request.json or {}
    name = data.get("name", "").strip()
    if len(name) < 2:
        return jsonify({"error": "团队名称至少2个字符"}), 400
    uid = session["user_id"]
    invite_code = secrets.token_hex(4)
    conn = get_user_db()
    cur = conn.execute("INSERT INTO teams (name, owner_id, invite_code) VALUES (?, ?, ?)",
                       (name, uid, invite_code))
    team_id = cur.lastrowid
    conn.execute("INSERT INTO team_members (team_id, user_id, role) VALUES (?, ?, 'owner')",
                 (team_id, uid))
    conn.commit()
    conn.close()
    return jsonify({"team_id": team_id, "invite_code": invite_code})


@app.route("/api/teams/<int:team_id>/join", methods=["POST"])
@login_required
def api_teams_join(team_id):
    """通过邀请码加入团队"""
    data = request.json or {}
    code = data.get("code", "").strip()
    uid = session["user_id"]
    conn = get_user_db()
    team = conn.execute("SELECT * FROM teams WHERE id=? AND invite_code=?", (team_id, code)).fetchone()
    if not team:
        conn.close()
        return jsonify({"error": "团队不存在或邀请码错误"}), 404
    existing = conn.execute("SELECT id FROM team_members WHERE team_id=? AND user_id=?",
                            (team_id, uid)).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": "你已在团队中"}), 400
    conn.execute("INSERT INTO team_members (team_id, user_id, role) VALUES (?, ?, 'member')",
                 (team_id, uid))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "team_name": team["name"]})


@app.route("/api/teams/<int:team_id>/members", methods=["GET"])
@login_required
def api_teams_members(team_id):
    """获取团队成员"""
    conn = get_user_db()
    members = conn.execute("""
        SELECT u.username, u.email, tm.role, tm.joined_at
        FROM team_members tm JOIN users u ON tm.user_id=u.id
        WHERE tm.team_id=? ORDER BY tm.role DESC, tm.joined_at ASC
    """, (team_id,)).fetchall()
    conn.close()
    return jsonify({"members": [dict(m) for m in members]})


@app.route("/api/teams/<int:team_id>/protocols", methods=["GET"])
@login_required
def api_teams_protocols(team_id):
    """获取团队共享Protocol"""
    uid = session["user_id"]
    conn = get_user_db()
    member = conn.execute("SELECT id FROM team_members WHERE team_id=? AND user_id=?",
                          (team_id, uid)).fetchone()
    if not member:
        conn.close()
        return jsonify({"error": "你不是该团队成员"}), 403
    protocols = conn.execute("""
        SELECT tp.protocol_id, tp.is_private, tp.created_at, u.username as added_by_name
        FROM team_protocols tp JOIN users u ON tp.added_by=u.id
        WHERE tp.team_id=?
        ORDER BY tp.created_at DESC LIMIT 50
    """, (team_id,)).fetchall()
    conn.close()
    return jsonify({"protocols": [dict(p) for p in protocols]})


@app.route("/api/teams/<int:team_id>/protocols", methods=["POST"])
@login_required
def api_teams_add_protocol(team_id):
    """向团队添加Protocol"""
    data = request.json or {}
    pid = data.get("protocol_id", "").strip()
    uid = session["user_id"]
    if not pid:
        return jsonify({"error": "请提供protocol_id"}), 400
    conn = get_user_db()
    try:
        conn.execute("INSERT INTO team_protocols (team_id, protocol_id, added_by) VALUES (?, ?, ?)",
                     (team_id, pid, uid))
        conn.commit()
    except:
        conn.close()
        return jsonify({"error": "添加失败，可能已存在"}), 400
    conn.close()
    return jsonify({"ok": True})


# ========== 管理后台 ==========

@app.route("/admin")
@admin_required
def admin_dashboard():
    conn = get_user_db()
    users = conn.execute("SELECT * FROM users ORDER BY created_at DESC").fetchall()
    total_users = conn.execute("SELECT COUNT(*) as c FROM users").fetchone()["c"]
    total_data = conn.execute("SELECT COUNT(*) as c FROM user_data").fetchone()["c"]
    conn.close()
    return render_template("admin.html", users=users, total_users=total_users, total_data=total_data)


@app.route("/admin/api/delete-user/<int:uid>", methods=["POST"])
@admin_required
def delete_user(uid):
    if uid == session.get("user_id"):
        return jsonify({"error": "不能删除自己"}), 400
    conn = get_user_db()
    conn.execute("DELETE FROM user_data WHERE user_id=?", (uid,))
    conn.execute("DELETE FROM users WHERE id=?", (uid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/admin/api/reset-password/<int:uid>", methods=["POST"])
@admin_required
def reset_password(uid):
    new_pw = secrets.token_hex(6)
    conn = get_user_db()
    conn.execute("UPDATE users SET password_hash=? WHERE id=?", (generate_password_hash(new_pw), uid))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "new_password": new_pw})


# ========== 静态文件路由（Vercel 兼容） ==========
@app.route("/static/<path:filename>")
def serve_static(filename):
    return send_from_directory(os.path.join(BASE_DIR, "static"), filename)

# ========== 路由 ==========

@app.route("/")
def home():
    """首页 - Protocol卡片网格 + 功能入口"""
    return render_template("home.html", protocols=PROTOCOL_META)


@app.route("/protocol/<protocol_id>")
def protocol_detail(protocol_id):
    """Protocol详情页 - 分tab展示"""
    meta = next((p for p in PROTOCOL_META if p["id"] == protocol_id), None)
    if not meta:
        return "Protocol not found", 404

    filepath = os.path.join(PROTOCOL_DIR, meta["file"])
    content = parse_protocol(filepath)

    # 获取关联仪器
    related_instruments = [i for i in INSTRUMENT_META if protocol_id in i.get("protocols", [])]

    return render_template("protocol.html", meta=meta, content=content,
                           instruments=related_instruments)


@app.route("/calculator")
def calculator():
    """试剂计算器页面"""
    return render_template("calculator.html", mw_db=MW_DB)


@app.route("/data")
def data_processing():
    """数据处理中心页面"""
    return render_template("data.html")


@app.route("/api/calculate", methods=["POST"])
def api_calculate():
    """计算器API"""
    data = request.json or {}
    calc_type = data.get("type")

    try:
        return _calc_impl(data, calc_type)
    except (KeyError, ValueError, TypeError, ZeroDivisionError) as e:
        return jsonify({"error": f"参数错误: {str(e)}"}), 400


def _calc_impl(data, calc_type):

    if calc_type == "molarity_to_mass":
        mw = float(data["mw"])
        molarity = float(data["molarity"])
        volume = float(data["volume"])
        mass_g = molarity * mw * (volume / 1000)
        return jsonify({
            "mass_g": round(mass_g, 4),
            "mass_mg": round(mass_g * 1000, 2),
            "volume": volume,
        })

    elif calc_type == "mass_to_molarity":
        mw = float(data["mw"])
        mass = float(data["mass"])
        volume = float(data["volume"])
        molarity = mass / (mw * volume / 1000)
        return jsonify({
            "molarity_mol": round(molarity, 6),
            "molarity_mmol": round(molarity * 1000, 4),
            "mass_conc": round(mass * 1000 / volume, 4),
        })

    elif calc_type == "dilution":
        stock = float(data["stock"])
        target = float(data["target"])
        volume = float(data["volume"])
        if stock < target:
            return jsonify({"error": "母液浓度必须大于目标浓度"})
        vol_stock = (target * volume) / stock
        vol_solvent = volume - vol_stock
        return jsonify({
            "vol_stock_ml": round(vol_stock, 4),
            "vol_stock_ul": round(vol_stock * 1000, 2),
            "vol_solvent": round(vol_solvent, 4),
            "dilution_ratio": round(stock / target, 2),
        })

    elif calc_type == "gradient":
        start = float(data["start"])
        factor = float(data["factor"])
        steps = int(data["steps"])
        vol = float(data["vol"])
        result = []
        for i in range(steps + 1):
            conc = start / (factor ** i)
            if i == 0:
                result.append({"level": f"第{i}级(原液)", "conc": round(conc, 6), "method": f"取原液 {vol}mL"})
            else:
                vs = (conc * vol) / start
                vx = vol - vs
                result.append({
                    "level": f"第{i}级(稀释{factor**i:.0f}倍)",
                    "conc": round(conc, 6),
                    "method": f"取原液 {round(vs*1000,2)}μL + 溶剂 {round(vx,3)}mL",
                })
        return jsonify({"table": result})

    elif calc_type == "rpm_rcf":
        mode = data["mode"]
        radius = float(data["radius"])
        if mode == "rpm_to_rcf":
            rpm = float(data["value"])
            rcf = 1.118e-5 * radius * rpm ** 2
            return jsonify({"result": round(rcf, 2), "unit": "×g"})
        else:
            rcf = float(data["value"])
            rpm = math.sqrt(rcf / (1.118e-5 * radius))
            return jsonify({"result": round(rpm, 0), "unit": "rpm"})

    elif calc_type == "reconstitution":
        mass_mg = float(data["mass"])
        conc_mg_ml = float(data["conc"])
        mw = float(data["mw"]) if data.get("mw") else 0
        volume_ml = mass_mg / conc_mg_ml if conc_mg_ml != 0 else 0
        result = {
            "volume_ml": round(volume_ml, 4),
            "volume_ul": round(volume_ml * 1000, 2),
        }
        if mw > 0:
            molarity = (mass_mg / (mw * 1000)) / (volume_ml / 1000) if volume_ml > 0 else 0
            result["molarity_mmol"] = round(molarity * 1000, 4)
        return jsonify(result)

    elif calc_type == "specific_activity":
        ed50 = float(data["ed50"])
        mw = float(data.get("mw", 0))
        if ed50 <= 0:
            return jsonify({"error": "ED50必须大于0"})
        sa = 1e6 / ed50
        result = {
            "sa_unit_mg": round(sa, 4),
        }
        if mw > 0:
            sa_mol = 1e6 / (ed50 * mw)
            result["sa_nmol_mg"] = round(sa_mol, 4)
        return jsonify(result)

    return jsonify({"error": "未知计算类型"})


@app.route("/instruments")
def instruments():
    """仪器指南列表页"""
    return render_template("instruments.html", instruments=INSTRUMENT_META, svg=get_instrument_svg)


@app.route("/instrument/<instrument_id>")
def instrument_detail(instrument_id):
    """仪器详情页 - 热区交互"""
    inst = next((i for i in INSTRUMENT_META if i["id"] == instrument_id), None)
    if not inst:
        return "Instrument not found", 404

    related_protocols = [p for p in PROTOCOL_META if p["id"] in inst.get("protocols", [])]
    svg_html = get_instrument_svg(inst.get("icon", ""))
    return render_template("instrument.html", inst=inst, protocols=related_protocols, svg=svg_html)


@app.route("/ai")
def ai_chat():
    """AI问答助手页面"""
    return render_template("ai.html")


def search_protocols_for_context(query, k=3):
    """基于关键词搜索Protocol，返回相关内容作为上下文"""
    keywords = query.lower().split()
    scores = []
    for meta in PROTOCOL_META:
        content = get_protocol_content(meta["id"])
        if not content:
            continue
        score = 0
        name_lower = meta["name"].lower()
        content_lower = content.lower()
        for kw in keywords:
            if kw in name_lower:
                score += 10
            if kw in content_lower:
                score += content_lower.count(kw)
        if score > 0:
            scores.append((score, meta, content))
    scores.sort(key=lambda x: x[0], reverse=True)
    results = []
    for _, meta, content in scores[:k]:
        results.append(f"【{meta['id']} {meta['name']}】\n{content[:2000]}")
    return "\n\n---\n\n".join(results)


@app.route("/api/chat", methods=["POST"])
def api_chat():
    """AI问答API - 基于本地Protocol知识库"""
    data = request.json
    user_msg = data.get("message", "")
    history = data.get("history", [])

    if not user_msg:
        return jsonify({"error": "请输入问题"})

    # 检索相关Protocol
    context = ""
    try:
        context = search_protocols_for_context(user_msg, k=3)
    except:
        pass

    answer = generate_local_answer(query=user_msg, context=context)
    return jsonify({"response": answer})


def generate_local_answer(query, context):
    """基于本地Protocol知识库的回答"""
    query_lower = query.lower()

    # 计算类问题
    if any(kw in query_lower for kw in ["算", "称", "配", "浓度", "稀释", "摩尔", "质量", "体积"]):
        answer = "### 🧪 试剂计算\n\n"
        answer += "请使用 **试剂计算器** 页面，支持以下功能：\n\n"
        answer += "| 计算类型 | 说明 |\n|---------|------|\n"
        answer += "| **稀释计算** | C₁V₁ = C₂V₂，支持mM/μM/mg/mL等单位 |\n"
        answer += "| **摩尔浓度** | 输入浓度+体积，自动算需称多少克 |\n"
        answer += "| **复溶计算** | 输入质量和目标浓度，算加多少溶剂 |\n"
        answer += "| **比活力计算** | 输入ED₅₀值，算Specific Activity |\n"
        answer += "| **单位换算** | 浓度、体积、rpm↔rcf互算 |\n\n"
        answer += "👉 [点击打开试剂计算器](/calculator)\n\n"
        if context:
            answer += "**相关Protocol配方参考：**\n\n"
            for line in context.split("\n")[:15]:
                if line.strip():
                    answer += f"> {line.strip()}\n"
        return answer

    # 失败排查
    if any(kw in query_lower for kw in ["失败", "没有", "不出", "不对", "偏低", "偏高", "怎么办", "异常", "问题"]):
        answer = "### 🔬 实验失败排查\n\n"
        answer += "实验失败很常见！帮你定位问题，请告诉我：\n\n"
        answer += "1. **你做的什么实验？**\n2. **具体现象是什么？**\n3. **操作细节？**\n\n"
        answer += "| 现象 | 可能原因 | 解决方案 |\n|------|---------|---------|\n"
        answer += "| 标准曲线R²<0.99 | 标准液配制不准 | 重新配制，每个点做3个重复 |\n"
        answer += "| OD值>0.8 | 样品浓度太高 | 稀释样品后重测 |\n"
        answer += "| 没有显色 | 试剂失效/温度不够 | 检查试剂有效期，确保沸水浴 |\n"
        answer += "| DNA条带模糊 | 降解/量太少 | 全程低温，增加样品量 |\n"
        answer += "| RNA降解 | RNase污染 | 全程RNase-free操作 |\n\n"
        if context:
            answer += "**相关Protocol避坑提示：**\n\n"
            for line in context.split("\n")[:15]:
                if line.strip() and ("做错" in line or "避坑" in line or "注意" in line or "不要" in line):
                    answer += f"> {line.strip()}\n"
        answer += "\n💡 也可以在 [Protocol库](/) 中查看对应实验的「避坑指南」Tab。"
        return answer

    # 有Protocol上下文时，给出结构化回答
    if context:
        answer = "### 📋 找到相关Protocol\n\n"

        # 解析protocol内容，提取关键信息
        lines = context.split("\n")
        current_section = ""
        sections = {}
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("一、") or line.startswith("二、") or line.startswith("三、") or \
               line.startswith("四、") or line.startswith("五、") or line.startswith("六、") or \
               line.startswith("七、") or line.startswith("八、") or line.startswith("九、") or \
               line.startswith("十、"):
                current_section = line
                sections[current_section] = []
            elif current_section and line:
                sections.setdefault(current_section, []).append(line)

        # 组织回答
        if "七、分步操作步骤" in sections or "七、分步操作" in context:
            for sec_name, sec_lines in sections.items():
                if "步骤" in sec_name or "操作" in sec_name:
                    answer += f"**{sec_name}**\n\n"
                    for l in sec_lines[:12]:
                        answer += f"- {l}\n"
                    answer += "\n"
                    break

        # 提取安全/避坑
        for sec_name, sec_lines in sections.items():
            if "安全" in sec_name or "避坑" in sec_name:
                answer += f"**{sec_name}**\n\n"
                for l in sec_lines[:8]:
                    answer += f"> {l}\n"
                answer += "\n"

        # 如果没有解析到结构化内容，给原始内容
        if answer == "### 📋 找到相关Protocol\n\n":
            for line in lines[:25]:
                if line.strip():
                    answer += f"> {line.strip()}\n"

        answer += "\n---\n\n"
        answer += "💡 在 [Protocol库](/) 中可以查看完整内容，每个步骤都有详细的操作指南。\n"
        answer += "如果你有更具体的问题，请告诉我！"
        return answer

    # 默认回答
    return """### 🌿 你好！我是植研小白盒AI助手

我可以帮你解答植物实验相关的问题：

**🧪 试剂配制**
- 帮你计算摩尔浓度、稀释倍数
- 告诉你具体称多少克、加多少溶剂

**🔬 实验操作**
- 解释实验原理
- 指导操作步骤
- 说明每步「做对了看到什么」

**❌ 失败排查**
- 分析实验失败原因
- 给出解决方案

**📊 数据处理**
- 标准曲线怎么做
- 数据怎么计算

👉 试试问我：「配100mL 0.1mol/L Tris需要称多少克？」"""


@app.route("/search")
def search_page():
    """Protocol检索页面"""
    return render_template("search.html")


@app.route("/api/search", methods=["POST"])
def api_search():
    """Protocol检索API - 全文关键词搜索 + 分类/难度筛选"""
    data = request.json
    query = data.get("query", "").strip()
    category = data.get("category", "")
    difficulty = data.get("difficulty", "")
    sort_by = data.get("sort", "relevance")

    keywords = query.split() if query else []

    conn = get_user_db()
    results = []
    for meta in PROTOCOL_META:
        # 分类筛选
        if category and meta["category"] != category:
            continue
        # 难度筛选
        if difficulty and str(meta["difficulty"]) != str(difficulty):
            continue

        content = get_protocol_content(meta["id"])
        if not content:
            continue

        score = 0
        content_lower = content.lower()
        name_lower = meta["name"].lower()
        desc_lower = meta.get("desc", "").lower()

        for kw in keywords:
            kw_lower = kw.lower()
            if kw_lower in name_lower:
                score += 10
            if kw_lower in desc_lower:
                score += 5
            # 内容匹配
            count = content_lower.count(kw_lower)
            score += count

        # 如果有筛选但无关键词，也返回全部匹配项
        if not keywords and (category or difficulty):
            score = 1

        if score > 0:
            snippet = ""
            for kw in keywords:
                idx = content_lower.find(kw.lower())
                if idx >= 0:
                    start = max(0, idx - 40)
                    end = min(len(content), idx + 80)
                    snippet = content[start:end].replace("\n", " ")
                    break

            # 获取收藏数和评分
            fav_cnt = conn.execute(
                "SELECT COUNT(*) as c FROM protocol_favorites WHERE protocol_id=?",
                (meta["id"],)
            ).fetchone()["c"]
            rating_row = conn.execute(
                "SELECT AVG(rating) as avg_r FROM protocol_ratings WHERE protocol_id=?",
                (meta["id"],)
            ).fetchone()
            avg_rating = round(rating_row["avg_r"], 1) if rating_row["avg_r"] else 0

            results.append({
                "protocol_id": meta["id"],
                "protocol_name": meta["name"],
                "category": meta["category"],
                "difficulty": meta["difficulty"],
                "desc": meta.get("desc", ""),
                "score": score,
                "snippet": snippet,
                "favorites": fav_cnt,
                "avg_rating": avg_rating,
            })

    conn.close()

    # 排序
    if sort_by == "favorites":
        results.sort(key=lambda x: x["favorites"], reverse=True)
    elif sort_by == "rating":
        results.sort(key=lambda x: x["avg_rating"], reverse=True)
    elif sort_by == "difficulty":
        results.sort(key=lambda x: x["difficulty"])
    else:
        results.sort(key=lambda x: x["score"], reverse=True)

    return jsonify({"results": results})


# ========== 导出功能 ==========
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill




# ========== AI数据分析 ==========

@app.route("/api/user/test-api-key", methods=["POST"])
@login_required
def api_user_test_api_key():
    """测试API Key是否可用"""
    data = request.json or {}
    provider = data.get("provider", "anthropic")
    api_key = data.get("api_key", "").strip()
    model = data.get("model", "").strip()
    base_url = data.get("base_url", "").strip()

    if not api_key:
        return jsonify({"error": "请先输入API Key"}), 400

    settings = {"provider": provider, "api_key": api_key, "model": model, "base_url": base_url}

    import urllib.request, urllib.error

    try:
        if provider == "anthropic":
            url = "https://api.anthropic.com/v1/messages"
            req_data = json.dumps({
                "model": model or "claude-sonnet-4-20250514",
                "max_tokens": 10,
                "messages": [{"role": "user", "content": "Hi"}],
            }).encode("utf-8")
            req = urllib.request.Request(url, data=req_data, headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            })
            with urllib.request.urlopen(req, timeout=30) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            model_used = result.get("model", model)
            return jsonify({"ok": True, "message": f"Anthropic连接成功 √ | 模型: {model_used}", "model": model_used})
        else:
            default_urls = {
                "deepseek": "https://api.deepseek.com/v1",
                "openai": "https://api.openai.com/v1",
            }
            api_base = base_url or default_urls.get(provider, "https://api.openai.com/v1")
            url = api_base.rstrip("/") + "/chat/completions"
            req_data = json.dumps({
                "model": model,
                "messages": [{"role": "user", "content": "Hi"}],
                "max_tokens": 10,
            }).encode("utf-8")
            req = urllib.request.Request(url, data=req_data, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            })
            with urllib.request.urlopen(req, timeout=30) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            model_used = result.get("model", model)
            tokens = result.get("usage", {}).get("total_tokens", 0)
            return jsonify({"ok": True, "message": f"{provider}连接成功 √ | 模型: {model_used} | 消耗{tokens}tokens", "model": model_used})

    except urllib.error.HTTPError as e:
        err_body = ""
        try:
            err_body = e.read().decode("utf-8")[:300]
        except:
            pass
        error_msg = f"连接失败: HTTP {e.code}"
        if e.code == 401:
            error_msg = "API Key无效 (401 Unauthorized)，请检查Key是否正确"
        elif e.code == 403:
            error_msg = "权限不足 (403 Forbidden)，请检查API Key权限或账户余额"
        elif e.code == 429:
            error_msg = "请求频率超限 (429)，请稍后重试"
        elif e.code == 404:
            error_msg = f"接口不存在 (404)，请检查Base URL是否正确: {url}"
        return jsonify({"error": error_msg, "detail": err_body}), 200

    except Exception as e:
        error_msg = str(e)
        if "Name or service not known" in error_msg or "getaddrinfo" in error_msg.lower():
            error_msg = f"无法解析域名，请检查Base URL: {base_url or default_urls.get(provider, '')}"
        elif "timed out" in error_msg.lower():
            error_msg = "连接超时，请检查网络"
        return jsonify({"error": f"连接异常: {error_msg}"}), 200


@app.route("/api/user/has-api-key")
def api_user_has_api_key():
    """检查用户是否配置了API Key"""
    settings = get_user_api_settings()
    return jsonify({
        "has_api_key": bool(settings and settings.get("api_key")),
        "provider": settings.get("provider") if settings else None,
    })


@app.route("/api/data/ai-analyze", methods=["POST"])
def api_data_ai_analyze():
    """AI智能分析数据：异步执行，返回task_id"""
    data = request.json or {}
    headers = data.get("headers", [])
    rows = data.get("rows", [])[:20]
    col_types = data.get("col_types", [])

    if not rows:
        return jsonify({"error": "没有数据可供分析"}), 400

    api_settings = get_user_api_settings()
    form_api_key = data.get("api_key", "")
    if form_api_key and not api_settings:
        api_settings = {"provider": "anthropic", "api_key": form_api_key, "model": "claude-sonnet-4-20250514", "base_url": ""}

    if not api_settings or not api_settings.get("api_key"):
        return jsonify({"error": "请先配置API Key", "need_api_key": True, "message": "AI分析需要调用大模型，请在个人中心配置API Key"})

    user_id = session.get("user_id")
    conn = get_user_db()
    cur = conn.execute(
        "INSERT INTO extract_tasks (user_id, task_type, status, file_name, raw_text) VALUES (?, 'ai_analyze', 'pending', 'data-analysis', ?)",
        (user_id, json.dumps({"headers": headers, "col_types": col_types, "row_count": len(rows)}, ensure_ascii=False))
    )
    task_id = cur.lastrowid
    conn.commit()
    conn.close()

    threading.Thread(target=_run_ai_analyze_task, args=(task_id, headers, rows, col_types, api_settings), daemon=True).start()
    return jsonify({"task_id": task_id, "status": "pending", "message": "AI正在分析..."})


def _build_data_summary(headers, rows, col_types):
    """构建数据摘要供AI分析"""
    parts = []
    parts.append(f"表头: {headers}")
    parts.append(f"总行数: {len(rows)}, 总列数: {len(headers)}")
    parts.append(f"列类型: {json.dumps([{'idx': c.get('index', i), 'header': c.get('header', ''), 'type': c.get('type', 'unknown')} for i, c in enumerate(col_types)], ensure_ascii=False)}")
    for ci in range(len(headers)):
        col_vals = []
        for row in rows:
            if ci < len(row):
                val = row[ci]
                if isinstance(val, (int, float)): col_vals.append(val)
        if col_vals:
            avg = sum(col_vals) / len(col_vals)
            srt = sorted(col_vals)
            parts.append(f"列{ci}[{headers[ci] if ci < len(headers) else ''}]: n={len(col_vals)}, mean={avg:.3f}, min={srt[0]:.3f}, max={srt[-1]:.3f}")
    parts.append(f"前{min(5, len(rows))}行:")
    for i, row in enumerate(rows[:5]):
        parts.append(f"  行{i}: {row}")
    return "\\n".join(parts)


def _run_ai_analyze_task(task_id, headers, rows, col_types, api_settings):
    try:
        conn = get_user_db()
        conn.execute("UPDATE extract_tasks SET status='processing' WHERE id=?", (task_id,))
        conn.commit()
        conn.close()

        data_summary = _build_data_summary(headers, rows, col_types)
        prompt = f"""植物科研数据分析专家。分析以下数据，返回JSON配置。

{data_summary}

返回JSON（不要markdown）:
{{"data_structure":"standard或group_indicator","experiment_type":"chlorophyll/soluble_protein/.../generic_stats/group_indicator_stats","confidence":0.8,"column_analysis":[{{"col_index":0,"suggested_type":"类型"}}],"column_mapping":{{"groupCol":0,"indicatorCol":1}},"suggested_params":{{}},"insights":[],"warnings":[],"recommended_actions":[]}}
列类型: sample_name/od_value/concentration/weight/group_col/indicator_col/numeric"""

        result_text, usage = call_llm(prompt, api_settings)
        structured = _extract_json_from_llm(result_text) if result_text else None

        conn = get_user_db()
        if structured:
            conn.execute("UPDATE extract_tasks SET status='completed', structured_json=?, result_json=?, completed_at=CURRENT_TIMESTAMP WHERE id=?", (json.dumps({"ai_analysis": structured}, ensure_ascii=False), json.dumps({"raw": result_text[:500] if result_text else ""}, ensure_ascii=False), task_id))
        else:
            conn.execute("UPDATE extract_tasks SET status='completed', structured_json=?, completed_at=CURRENT_TIMESTAMP WHERE id=?", (json.dumps({"error": "AI解析失败"}, ensure_ascii=False), task_id))
        conn.commit()
        conn.close()
    except Exception as e:
        conn = get_user_db()
        conn.execute("UPDATE extract_tasks SET status='failed', error_msg=?, completed_at=CURRENT_TIMESTAMP WHERE id=?", (str(e)[:500], task_id))
        conn.commit()
        conn.close()


@app.route("/api/data/ai-analyze/status/<int:task_id>")
def api_ai_analyze_status(task_id):
    conn = get_user_db()
    row = conn.execute("SELECT * FROM extract_tasks WHERE id=? AND task_type='ai_analyze'", (task_id,)).fetchone()
    conn.close()
    if not row: return jsonify({"error": "任务不存在"}), 404
    result = {"task_id": row["id"], "status": row["status"]}
    if row["status"] == "completed": result["ai_analysis"] = json.loads(row["structured_json"]) if row["structured_json"] else {}
    elif row["status"] == "failed": result["error"] = row["error_msg"]
    return jsonify(result)

@app.route("/api/export-excel", methods=["POST"])
def export_excel():
    """导出数据处理结果为Excel格式"""
    data = request.json
    sheet_name = data.get("sheetName", "数据处理结果")
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    title = data.get("title", "")

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # 字体定义：中文宋体+英文Times New Roman，5号(10.5pt)
    font_header = Font(name="宋体", size=10.5, bold=True)
    font_data = Font(name="宋体", size=10.5)
    font_title = Font(name="宋体", size=14, bold=True)

    # 对齐
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # 边框
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # 标题行
    start_row = 1
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        cell = ws.cell(row=1, column=1, value=title)
        cell.font = font_title
        cell.alignment = align_center
        start_row = 2

    # 表头
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=h)
        cell.font = font_header
        cell.alignment = align_center
        cell.border = thin_border
        cell.fill = PatternFill(start_color="D9E2C8", end_color="D9E2C8", fill_type="solid")

    # 数据行
    for row_idx, row_data in enumerate(rows, start_row + 1):
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = font_data
            cell.alignment = align_center
            cell.border = thin_border

    # 自动列宽（根据内容计算）
    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                # 中文字符占2个宽度
                val = str(cell.value)
                length = sum(2 if ord(c) > 127 else 1 for c in val)
                max_len = max(max_len, length)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)

    # 自动行高（根据内容自动扩展）
    for row in ws.rows:
        max_lines = 1
        for cell in row:
            if cell.value:
                val = str(cell.value)
                # 计算需要的行数（根据列宽估算）
                col_width = ws.column_dimensions[cell.column_letter].width or 10
                chars_per_line = max(col_width / 2, 5)
                lines = math.ceil(len(val) / chars_per_line) if len(val) > chars_per_line else 1
                max_lines = max(max_lines, lines)
        ws.row_dimensions[row[0].row].height = max(max_lines * 15, 20)

    # 输出
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"{sheet_name}.xlsx",
    )


@app.route("/export")
def export_page():
    return render_template("export.html", protocols=PROTOCOL_META)

@app.route("/api/export/<pid>/<fmt>")
def export_protocol(pid, fmt):
    """导出Protocol为md/docx"""
    matching = None
    for p in PROTOCOL_META:
        if p["id"] == pid:
            matching = p
            break
    if not matching:
        return jsonify({"error": "Protocol not found"}), 404

    filepath = os.path.join(PROTOCOL_DIR, matching["file"])
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    if fmt == "md":
        md_content = convert_to_markdown(content, matching)
        buf = io.BytesIO()
        buf.write(md_content.encode("utf-8"))
        buf.seek(0)
        return send_file(buf, mimetype="text/markdown",
                        as_attachment=True,
                        download_name=f"{pid}_{matching['name']}.md")

    elif fmt == "docx":
        doc = create_docx(content, matching)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        as_attachment=True,
                        download_name=f"{pid}_{matching['name']}.docx")

    return jsonify({"error": "Unsupported format"}), 400

def set_run_font(run, size_pt=10.5, bold=False):
    """设置run字体：中文宋体，英文Times New Roman，五号=10.5pt"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

def add_paragraph(doc, text, size=10.5, bold=False, align=None, space_after=Pt(3)):
    """添加一个段落，设置中英文混排字体"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size, bold)
    return p

def add_checkbox_line(doc, text, size=10.5):
    """添加带打勾空格的行：☐ text"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    # 空格方框 ☐
    run = p.add_run("☐ ")
    set_run_font(run, size)
    run = p.add_run(text)
    set_run_font(run, size)
    return p

def convert_to_markdown(content, meta):
    """将Protocol txt转为格式化Markdown"""
    lines = content.split("\n")
    md = []
    md.append(f"# {meta['id']} {meta['name']}\n")

    for line in lines:
        line = line.strip()
        if not line:
            md.append("")
            continue
        if re.match(r"^[一二三四五六七八九十]+[、.]", line):
            md.append(f"\n## {line}\n")
        elif re.match(r"^步骤\d+", line):
            md.append(f"\n### {line}\n")
        elif re.match(r"^\d+\.", line):
            md.append(f"- [ ] {line[2:]}")
        else:
            md.append(line)

    return "\n".join(md)

def create_docx(content, meta):
    """将Protocol txt转为Word文档 - 宋体五号/Times New Roman五号"""
    doc = DocxDocument()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 设置Normal样式
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 大标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(f"{meta['id']}  {meta['name']}")
    set_run_font(run, 16, bold=True)

    # ===== 来源出处 =====
    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_p.paragraph_format.space_after = Pt(12)
    run = src_p.add_run("植研小白盒 · 标准化Protocol")
    set_run_font(run, 9)
    run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    # ===== 解析内容 =====
    lines = content.split("\n")
    current_section = ""

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line:
            continue

        # 跳过Protocol编号和实验名称行（已在标题中）
        if line.startswith("Protocol编号") or line.startswith("实验名称"):
            continue

        # 一级标题：一、原理  二、材料  三、步骤 ...
        if re.match(r"^[一二三四五六七八九十]+[、.]", line):
            current_section = line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            # 标题下加横线
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): '2E7D32',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            run = p.add_run(line)
            set_run_font(run, 12, bold=True)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            continue

        # 步骤标题：步骤1：xxx
        if re.match(r"^步骤\d+", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("🔹 " + line)
            set_run_font(run, 11, bold=True)
            continue

        # 子项：为什么/做对了/做错了 -> 带缩进
        if line.startswith("为什么：") or line.startswith("做对了：") or line.startswith("做错了："):
            add_checkbox_line(doc, line)
            continue

        # 普通列表项 1. 2. 3.
        if re.match(r"^\d+\.", line):
            add_checkbox_line(doc, line)
            continue

        # 普通正文
        add_paragraph(doc, line)

    # ===== 底部签字栏 =====
    doc.add_paragraph()  # 空行
    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(20)
    run = sign_p.add_run("实验日期：____________    操作人：____________    复核人：____________")
    set_run_font(run, 10.5)

    return doc

# ========== 文献提取 ==========

# OpenAI兼容SDK（可选）
try:
    import openai as openai_sdk
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False


def call_llm(prompt, settings):
    """通用LLM调用，返回 (text, token_usage)"""
    import ssl
    provider = settings.get("provider", "anthropic")
    api_key = settings.get("api_key", "")
    model = settings.get("model", "")
    base_url = settings.get("base_url", "")

    if not api_key:
        return None, {}

    usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

    # Python 3.6 SSL兼容：创建不验证证书的context以兼容旧版OpenSSL
    ssl_context = None
    try:
        ssl_context = ssl.create_default_context()
    except:
        try:
            ssl_context = ssl._create_unverified_context()
        except:
            pass

    try:
        import urllib.request
        import urllib.error

        if provider == "anthropic":
            url = "https://api.anthropic.com/v1/messages"
            req_data = json.dumps({
                "model": model or "claude-sonnet-4-20250514",
                "max_tokens": 8192,
                "messages": [{"role": "user", "content": prompt}],
            }).encode("utf-8")
            req = urllib.request.Request(url, data=req_data, headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            })
            try:
                with urllib.request.urlopen(req, timeout=120, context=ssl_context) as resp:
                    result = json.loads(resp.read().decode("utf-8"))
                if "usage" in result:
                    usage = {
                        "prompt_tokens": result["usage"].get("input_tokens", 0),
                        "completion_tokens": result["usage"].get("output_tokens", 0),
                        "total_tokens": result["usage"].get("input_tokens", 0) + result["usage"].get("output_tokens", 0),
                    }
                return result["content"][0]["text"], usage
            except urllib.error.HTTPError as e:
                err_body = ""
                try: err_body = e.read().decode("utf-8")[:300]
                except: pass
                print(f"[ERROR] Anthropic API错误 ({model}): HTTP {e.code}")
                return None, usage
            except Exception as e:
                print(f"[ERROR] Anthropic请求异常: {e}")
                return None, usage
        else:
            default_urls = {
                "deepseek": "https://api.deepseek.com/v1",
                "openai": "https://api.openai.com/v1",
            }
            api_base = base_url or default_urls.get(provider, "https://api.openai.com/v1")
            url = api_base.rstrip("/") + "/chat/completions"

            req_data = json.dumps({
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 8192,
            }).encode("utf-8")

            req = urllib.request.Request(url, data=req_data, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            })

            try:
                with urllib.request.urlopen(req, timeout=120, context=ssl_context) as resp:
                    result = json.loads(resp.read().decode("utf-8"))
                if "usage" in result:
                    usage = {
                        "prompt_tokens": result["usage"].get("prompt_tokens", 0),
                        "completion_tokens": result["usage"].get("completion_tokens", 0),
                        "total_tokens": result["usage"].get("total_tokens", 0),
                    }
                return result["choices"][0]["message"]["content"], usage
            except urllib.error.HTTPError as e:
                err_body = ""
                try: err_body = e.read().decode("utf-8")[:300]
                except: pass
                print(f"[ERROR] API HTTP错误 ({provider}/{model}): {e.code}")
                return None, usage
            except Exception as e:
                print(f"[ERROR] API请求异常 ({provider}): {e}")
                return None, usage
    except Exception as e:
        print(f"[ERROR] LLM调用失败 ({provider}/{model}): {e}")
        return None, usage


def get_user_api_settings(user_id=None):
    """获取用户API设置，返回settings dict"""
    uid = user_id or session.get("user_id")
    if not uid:
        return None
    conn = get_user_db()
    row = conn.execute("SELECT * FROM user_api_settings WHERE user_id=?", (uid,)).fetchone()
    conn.close()
    if row and row["api_key"]:
        return {
            "provider": row["provider"],
            "api_key": row["api_key"],
            "model": row["model"],
            "base_url": row["base_url"] or "",
        }
    return None


@app.route("/extract")
def extract_page():
    """文献实验过程提取页面"""
    has_api = False
    if "user_id" in session:
        settings = get_user_api_settings()
        has_api = bool(settings and settings.get("api_key"))
    return render_template("extract.html", has_api=has_api)


def extract_structured_sections(text):
    """从论文全文中提取结构化实验内容 - 返回新格式"""
    # 找到"材料与方法"部分
    text_lower = text.lower()
    method_start = 0
    method_markers = ["材料与方法", "材料和方法", "实验方法", "实验材料与方法",
                       "materials and methods", "materials & methods",
                       "experimental procedures", "methods and materials"]
    for marker in method_markers:
        idx = text_lower.find(marker.lower())
        if idx >= 0:
            method_start = idx
            break

    # 找到方法部分的结束位置
    method_end = len(text)
    if method_start > 0:
        end_patterns = [
            r"\n(?:结果|讨论|结论|参考文献|致谢|acknowledgment|discussion|results|references)",
        ]
        for pat in end_patterns:
            m = re.search(pat, text[method_start + 10:], re.IGNORECASE)
            if m:
                method_end = method_start + 10 + m.start()
                break

    method_text = text[method_start:method_end]

    # 提取标题（方法部分之前的文本）
    title = ""
    before_method = text[:method_start] if method_start > 0 else text[:200]
    for line in before_method.split("\n"):
        line = line.strip()
        if line and len(line) > 5:
            title = line[:100]
            break

    # 按行解析方法部分
    lines = method_text.split("\n")

    # 识别子标题（如 "1. 实验材料" "2. RNA提取"）
    sub_sections = []
    current_section = None
    section_buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 检测子标题：数字开头 + 中文标题
        sub_title_match = re.match(r"^(\d+)[.、．]\s*(.{2,30})", stripped)
        if sub_title_match:
            if current_section:
                sub_sections.append({"title": current_section, "lines": section_buffer})
            current_section = sub_title_match.group(2).strip()
            section_buffer = []
            continue

        if current_section:
            section_buffer.append(stripped)

    if current_section:
        sub_sections.append({"title": current_section, "lines": section_buffer})

    # 如果没有子标题，把整个方法文本当作一个section
    if not sub_sections:
        all_lines = [l.strip() for l in lines if l.strip()]
        if all_lines:
            sub_sections.append({"title": "方法", "lines": all_lines})

    # 分类各子section
    result = {
        "title": title,
        "method_name": "",
        "principle": "",
        "steps": [],
        "reagents": [],
        "instruments": [],
        "conditions": {},
        "risks": [],
    }

    material_kw = ["材料", "试剂", "reagent", "material", "药品", "溶液"]
    instrument_kw = ["仪器", "设备", "instrument", "apparatus", "equipment"]
    condition_kw = ["条件", "condition", "培养", "生长条件"]
    note_kw = ["注意", "安全", "note", "warning", "caution", "避坑"]
    step_kw = ["步骤", "操作", "procedure", "protocol", "step", "方法", "测定", "提取", "检测", "分析"]

    for sec in sub_sections:
        sec_title_lower = sec["title"].lower()
        matched = None

        for kw in material_kw:
            if kw in sec_title_lower:
                matched = "reagents"
                break
        if not matched:
            for kw in instrument_kw:
                if kw in sec_title_lower:
                    matched = "instruments"
                    break
        if not matched:
            for kw in condition_kw:
                if kw in sec_title_lower:
                    matched = "conditions"
                    break
        if not matched:
            for kw in note_kw:
                if kw in sec_title_lower:
                    matched = "risks"
                    break
        if not matched:
            for kw in step_kw:
                if kw in sec_title_lower:
                    matched = "steps"
                    break

        # 如果标题没匹配上，用内容判断
        if not matched:
            content = " ".join(sec["lines"])
            if any(kw in content for kw in ["mg", "mL", "μL", "g/L", "mol/L", "mM", "浓度", "配制"]):
                matched = "reagents"
            elif any(kw in content for kw in ["rpm", "离心", "℃", "°C", "温度", "培养"]):
                matched = "conditions"
            elif any(kw in content for kw in ["仪器", "设备", "型号"]):
                matched = "instruments"
            else:
                matched = "steps"

        if matched == "reagents":
            for line in sec["lines"]:
                if len(line) < 3:
                    continue
                conc = ""
                cm = re.search(r"(\d+\.?\d*\s*(?:mol/L|mM|μM|mg/mL|g/L|%|mol·L))", line)
                if cm:
                    conc = cm.group(1)
                result["reagents"].append({
                    "name": line[:80],
                    "concentration": conc,
                    "amount": "",
                    "prep": "",
                    "source_confidence": "部分",
                })

        elif matched == "steps":
            for line in sec["lines"]:
                if len(line) < 5:
                    continue
                params = ""
                # 提取数字参数
                pm = re.search(r"(\d+\.?\d*\s*(?:rpm|°C|℃|min|h|μL|mL|g|mg))", line, re.IGNORECASE)
                if pm:
                    params = pm.group(1)
                result["steps"].append({
                    "step": len(result["steps"]) + 1,
                    "action": line,
                    "params": params,
                    "reagent": "",
                    "warning": "",
                    "source_confidence": "部分",
                })

        elif matched == "instruments":
            for line in sec["lines"]:
                if len(line) > 2:
                    # 按逗号或顿号分割
                    for inst in re.split(r"[,，、;；]", line):
                        inst = inst.strip()
                        if inst and len(inst) > 1:
                            result["instruments"].append(inst[:40])

        elif matched == "conditions":
            cond_text = " ".join(sec["lines"])
            temp_m = re.search(r"(\d+)\s*[°℃]", cond_text)
            if temp_m:
                result["conditions"]["temperature"] = temp_m.group(0)
            if "光" in cond_text:
                light_m = re.search(r"(\d+\s*h[^。]*光[^。]*)", cond_text)
                result["conditions"]["light"] = light_m.group(1)[:60] if light_m else cond_text[:60]
            if not result["conditions"]:
                result["conditions"]["other"] = cond_text[:100]

        elif matched == "risks":
            for line in sec["lines"]:
                if len(line) > 5:
                    result["risks"].append({
                        "point": line[:120],
                        "solution": "",
                        "severity": "中",
                    })

    # 如果没有提取到步骤，尝试从全文提取带数字的行
    if not result["steps"]:
        step_num = 0
        for line in lines:
            stripped = line.strip()
            if not stripped or len(stripped) < 5:
                continue
            # 匹配 "取..." "加..." "离心..." "混合..." 等操作动词开头的行
            if re.match(r"^(取|加|加入|混|离心|称|配|溶|稀释|测定|检测|提取|分离|纯化|孵育|静置|震荡|摇|过滤|洗涤|晾干|溶解|定容|滴加|吸取|转移|置于|放入|打开|关闭|设置)", stripped):
                step_num += 1
                pm = re.search(r"(\d+\.?\d*\s*(?:rpm|°C|℃|min|h|μL|mL|g|mg))", stripped, re.IGNORECASE)
                result["steps"].append({
                    "step": step_num,
                    "action": stripped,
                    "params": pm.group(1) if pm else "",
                    "reagent": "",
                    "warning": "",
                    "source_confidence": "部分",
                })

    # 从全文提取仪器（如果还没找到）
    if not result["instruments"]:
        for line in lines:
            if "仪器" in line or "设备" in line:
                for inst in re.split(r"[,，、;；]", line):
                    inst = inst.strip()
                    if inst and len(inst) > 1 and "仪器" not in inst and "设备" not in inst:
                        result["instruments"].append(inst[:40])

    return result


def extract_text_with_mineru(pdf_bytes):
    """使用MinerU解析PDF，返回高质量Markdown文本（支持表格、公式、多栏排版）"""
    import tempfile
    import os
    import glob as _glob

    with tempfile.TemporaryDirectory() as tmp_dir:
        mineru_do_parse(
            output_dir=tmp_dir,
            pdf_file_names=["paper.pdf"],
            pdf_bytes_list=[pdf_bytes],
            p_lang_list=["ch"],
            backend="pipeline",
            parse_method="auto",
            formula_enable=True,
            table_enable=True,
            f_draw_layout_bbox=False,
            f_draw_span_bbox=False,
            f_dump_md=True,
            f_dump_middle_json=False,
            f_dump_model_output=False,
            f_dump_orig_pdf=False,
            f_dump_content_list=False,
        )
        # 读取生成的markdown文件
        md_files = _glob.glob(os.path.join(tmp_dir, "**", "*.md"), recursive=True)
        if md_files:
            with open(md_files[0], "r", encoding="utf-8") as f:
                return f.read()
    return ""


@app.route("/api/extract", methods=["POST"])
def api_extract():
    """PDF文献提取API - 后台异步执行"""
    if not HAS_PDF and not HAS_MINERU:
        return jsonify({"error": "PDF解析库未安装，请联系管理员"}), 500

    if "file" not in request.files:
        return jsonify({"error": "请上传文件"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "请选择文件"}), 400

    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "仅支持PDF格式"}), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > 20 * 1024 * 1024:
        return jsonify({"error": "文件大小不能超过20MB"}), 400

    try:
        pdf_bytes = file.read()
        full_text = ""

        # 优先使用MinerU（高质量解析，支持表格、公式、多栏排版）
        if HAS_MINERU:
            try:
                full_text = extract_text_with_mineru(pdf_bytes)
            except Exception as e:
                print(f"[WARN] MinerU解析失败，回退到pdfplumber: {e}")

        # MinerU不可用或失败时，回退到pdfplumber
        if not full_text.strip() and HAS_PDF:
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                cells = [str(c).strip() if c else "" for c in row]
                                text_parts.append(" | ".join(cells))
            full_text = "\n".join(text_parts)

        if not full_text.strip():
            return jsonify({"error": "无法从PDF中提取文本，可能是扫描件。请使用文本型PDF。"}), 400

        api_settings = None
        if "user_id" in session:
            api_settings = get_user_api_settings()
        form_api_key = request.form.get("api_key", "")
        if form_api_key and not api_settings:
            api_settings = {"provider": "anthropic", "api_key": form_api_key, "model": "claude-sonnet-4-20250514", "base_url": ""}

        # 无API时同步返回（关键词模式很快）
        if not api_settings or not api_settings.get("api_key"):
            structured, usage = extract_structured_sections(full_text), {}
            return jsonify({
                "text": full_text, "structured": structured,
                "pages": len(text_parts), "ai_powered": False, "usage": usage,
            })

        # 有API时后台异步执行
        user_id = session.get("user_id")
        file_data_b64 = base64.b64encode(full_text.encode("utf-8")).decode("utf-8")
        conn = get_user_db()
        cur = conn.execute(
            "INSERT INTO extract_tasks (user_id, task_type, status, file_name, file_data, raw_text) VALUES (?, 'extract', 'pending', ?, ?, ?)",
            (user_id, file.filename, file_data_b64, full_text[:50000])
        )
        task_id = cur.lastrowid
        conn.commit()
        conn.close()

        threading.Thread(
            target=_run_extract_task,
            args=(task_id, api_settings),
            daemon=True
        ).start()

        return jsonify({"task_id": task_id, "status": "pending"})

    except Exception as e:
        return jsonify({"error": f"解析失败: {str(e)}"}), 500


def _run_extract_task(task_id, api_settings):
    """后台执行提取任务"""
    try:
        conn = get_user_db()
        conn.execute("UPDATE extract_tasks SET status='processing' WHERE id=?", (task_id,))
        conn.commit()

        row = conn.execute("SELECT file_data, raw_text FROM extract_tasks WHERE id=?", (task_id,)).fetchone()
        conn.close()

        full_text = base64.b64decode(row["file_data"]).decode("utf-8")
        structured, usage = extract_with_llm(full_text, api_settings)

        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='completed', structured_json=?, result_json=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (json.dumps(structured, ensure_ascii=False), json.dumps({"usage": usage}, ensure_ascii=False), task_id)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='failed', error_msg=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (str(e), task_id)
        )
        conn.commit()
        conn.close()


@app.route("/api/extract/task/<int:task_id>")
@login_required
def api_extract_task_status(task_id):
    """查询提取任务状态"""
    conn = get_user_db()
    row = conn.execute("SELECT * FROM extract_tasks WHERE id=? AND user_id=?", (task_id, session["user_id"])).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "任务不存在"}), 404
    result = {
        "task_id": row["id"],
        "task_type": row["task_type"],
        "status": row["status"],
        "file_name": row["file_name"],
        "created_at": row["created_at"],
        "completed_at": row["completed_at"],
    }
    if row["status"] == "completed":
        result["structured"] = json.loads(row["structured_json"]) if row["structured_json"] else {}
        result["raw_text"] = row["raw_text"] or ""
        result["usage"] = json.loads(row["result_json"]) if row["result_json"] else {}
        if row["issues_json"]:
            result["issues"] = json.loads(row["issues_json"])
    elif row["status"] == "failed":
        result["error"] = row["error_msg"]
    return jsonify(result)


@app.route("/api/extract/tasks")
@login_required
def api_extract_tasks():
    """获取用户的提取任务列表"""
    conn = get_user_db()
    rows = conn.execute(
        "SELECT id, task_type, status, file_name, created_at, completed_at FROM extract_tasks WHERE user_id=? ORDER BY id DESC LIMIT 20",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    return jsonify({"tasks": [
        {"task_id": r["id"], "task_type": r["task_type"], "status": r["status"],
         "file_name": r["file_name"], "created_at": r["created_at"], "completed_at": r["completed_at"]}
        for r in rows
    ]})


def _extract_json_from_llm(text):
    """从LLM输出中健壮地提取JSON，处理markdown代码块、多余文字等"""
    if not text:
        return None
    # 先尝试提取 ```json ... ``` 代码块
    m = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass
    # 尝试找最外层 { ... }
    start = text.find("{")
    if start < 0:
        return None
    depth = 0
    in_string = False
    escape = False
    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if c == '\\' and in_string:
            escape = True
            continue
        if c == '"' and not escape:
            in_string = not in_string
            continue
        if in_string:
            continue
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                try:
                    return json.loads(text[start:i+1])
                except json.JSONDecodeError:
                    # 尝试修复常见问题：尾逗号
                    candidate = text[start:i+1]
                    candidate = re.sub(r',\s*([}\]])', r'\1', candidate)
                    try:
                        return json.loads(candidate)
                    except json.JSONDecodeError:
                        return None
    return None


def _preprocess_methods_section(text):
    """预处理PDF文本，只保留材料与方法部分，减少token消耗"""
    text_lower = text.lower()

    # 找到方法部分的起始位置
    method_start = -1
    markers = ["材料与方法", "材料和方法", "实验方法", "实验材料与方法",
                "materials and methods", "materials & methods",
                "experimental procedures", "methods and materials",
                "materials and methods"]
    for marker in markers:
        idx = text_lower.find(marker.lower())
        if idx >= 0:
            method_start = idx
            break

    if method_start < 0:
        # 没找到方法部分，尝试用关键词密度判断
        # 找包含最多实验关键词的连续段落
        return text[:30000] if len(text) > 30000 else text

    # 找到方法部分的结束位置
    method_end = len(text)
    end_markers = [
        r"\n(?:结果|讨论|结论|参考文献|致谢|acknowledgment|discussion|results|references|图表|figure|table)",
        r"\n\d+\.\s*(?:结果|讨论|结论|Result|Discussion|Conclusion)",
    ]
    for pat in end_markers:
        m = re.search(pat, text[method_start:], re.IGNORECASE)
        if m:
            method_end = method_start + m.start()
            break

    # 提取标题（方法部分之前的文本，取前200字）
    title_area = text[:method_start][:300]

    # 组合：标题 + 方法部分
    methods_text = text[method_start:method_end]
    combined = title_area + "\n\n" + methods_text

    return combined


def extract_with_llm(text, api_settings):
    """用LLM API精准提取实验方法、步骤、试剂、参数"""
    # 预处理：只保留方法部分，减少token消耗
    text = _preprocess_methods_section(text)
    if len(text) > 30000:
        text = text[:30000] + "\n\n[...文本过长，已截断...]"

    prompt = f"""你是一个专业的植物科研实验助手。你的任务是从学术论文中精准提取"材料与方法"（Materials and Methods）部分，将其转化为一份**可直接执行的实验方案**。

请仔细阅读论文全文，找到"材料与方法"/"Materials and Methods"/"实验方法"部分，然后提取以下信息：

**要求：**
1. 步骤必须是**可执行的**——一个从未做过这个实验的人看了就能动手操作
2. 每个步骤必须包含：具体操作 + 关键参数（温度、时间、转速、体积、质量等）
3. 试剂要列出完整配方（名称、浓度、用量、配制方法）
4. 标注每个信息的来源可信度

**返回JSON格式（严格按此格式）：**
```json
{{
  "title": "论文完整标题",
  "method_name": "实验方法名称（如：TRIzol法提取总RNA）",
  "principle": "实验原理（1-2句话说明这个方法的原理）",
  "steps": [
    {{
      "step": 1,
      "action": "详细操作描述，包含所有动作和参数。例如：取0.1g新鲜叶片放入1.5mL离心管中，加入1mL TRIzol试剂，用液氮研磨至粉末状",
      "params": "关键参数摘要，如：0.1g, 1mL, 液氮",
      "reagent": "本步骤使用的试剂名称",
      "warning": "本步骤的注意事项（没有则留空字符串）",
      "source_confidence": "明确"
    }}
  ],
  "reagents": [
    {{
      "name": "试剂全称",
      "concentration": "工作浓度",
      "amount": "用量（如：1mL/样品）",
      "prep": "配制方法（如：TRIzol直接使用，无需配制）",
      "source_confidence": "明确"
    }}
  ],
  "instruments": ["仪器名称及型号（如有）"],
  "conditions": {{
    "temperature": "温度条件",
    "light": "光照条件（如有）",
    "other": "其他重要条件"
  }},
  "risks": [
    {{
      "point": "具体风险描述",
      "solution": "解决方案或替代方法",
      "severity": "高/中/低"
    }}
  ]
}}
```

**source_confidence 判断标准：**
- "明确"：论文中**明确写了**这个数值/操作（如"12000rpm离心10min"）
- "部分"：论文提到了但**参数不完整**（如只写了"离心"但没写转速）
- "推断"：论文**没有明确写**，你根据实验常识推断的（如"室温"推断为25°C）
- "缺失"：论文**完全没有提及**这个信息

**风险点（risks）应包括：**
- 论文中明确提到的注意事项
- 你根据实验经验判断的常见出错点
- 关键步骤中容易忽略的操作细节

**重要原则：**
- 只提取论文中**实际存在的**实验操作，不要编造
- 步骤要**完整、详细、按顺序**，不要遗漏
- 每个参数都要**保留原文数值**，不要修改
- 用中文返回（术语可保留英文）
- 严格返回JSON，不要有其他文字

论文内容：
{text}"""

    try:
        result_text, usage = call_llm(prompt, api_settings)
        if not result_text:
            print("[WARN] LLM返回空结果，回退到关键词模式")
            return extract_structured_sections(text), {}

        structured = _extract_json_from_llm(result_text)
        if structured:
            if "steps" in structured and isinstance(structured["steps"], str):
                structured = _convert_old_format(structured)
            return structured, usage
        else:
            print(f"[WARN] LLM返回的JSON解析失败，回退到关键词模式。原文前200字: {result_text[:200]}")
            return extract_structured_sections(text), {}

    except Exception as e:
        print(f"[WARN] LLM API提取失败，回退到关键词模式: {e}")
        return extract_structured_sections(text), {}


def _convert_old_format(data):
    """将旧格式（steps为字符串）转换为新格式（steps为数组）"""
    result = {
        "title": data.get("title", ""),
        "method_name": data.get("methods", ""),
        "principle": "",
        "steps": [],
        "reagents": [],
        "instruments": [i.strip() for i in data.get("instruments", "").split(",") if i.strip()] if data.get("instruments") else [],
        "conditions": {},
        "risks": [],
    }
    # 解析旧格式的steps字符串
    steps_text = data.get("steps", "")
    if isinstance(steps_text, str):
        for i, line in enumerate(steps_text.split("\n"), 1):
            line = line.strip()
            if line and len(line) > 5:
                result["steps"].append({
                    "step": i,
                    "action": line,
                    "params": "",
                    "reagent": "",
                    "warning": "",
                    "source_confidence": "部分",
                })
    # 解析旧格式的materials
    mats = data.get("materials", "")
    if isinstance(mats, str) and mats.strip():
        for line in mats.split("\n"):
            line = line.strip()
            if line and len(line) > 3:
                result["reagents"].append({
                    "name": line,
                    "concentration": "",
                    "amount": "",
                    "prep": "",
                    "source_confidence": "部分",
                })
    if data.get("conditions"):
        result["conditions"] = {"other": data["conditions"]}
    if data.get("notes"):
        result["risks"].append({"point": data["notes"], "solution": "", "severity": "中"})
    return result


@app.route("/api/extract/review", methods=["POST"])
@login_required
def api_extract_review():
    """AI审查提取结果 - 后台异步执行"""
    data = request.json
    structured = data.get("structured", {})
    raw_text = data.get("raw_text", "")

    api_settings = get_user_api_settings()
    form_api_key = data.get("api_key", "")
    if form_api_key and not api_settings:
        api_settings = {"provider": "anthropic", "api_key": form_api_key, "model": "claude-sonnet-4-20250514", "base_url": ""}

    if not api_settings or not api_settings.get("api_key"):
        return jsonify({"error": "请先在个人中心配置API Key"})

    user_id = session["user_id"]
    conn = get_user_db()
    cur = conn.execute(
        "INSERT INTO extract_tasks (user_id, task_type, status, structured_json, raw_text) VALUES (?, 'review', 'pending', ?, ?)",
        (user_id, json.dumps(structured, ensure_ascii=False), raw_text[:50000])
    )
    task_id = cur.lastrowid
    conn.commit()
    conn.close()

    threading.Thread(
        target=_run_review_task,
        args=(task_id, api_settings),
        daemon=True
    ).start()

    return jsonify({"task_id": task_id, "status": "pending"})


def _run_review_task(task_id, api_settings):
    """后台执行审查任务"""
    try:
        conn = get_user_db()
        conn.execute("UPDATE extract_tasks SET status='processing' WHERE id=?", (task_id,))
        conn.commit()
        row = conn.execute("SELECT structured_json, raw_text FROM extract_tasks WHERE id=?", (task_id,)).fetchone()
        conn.close()

        structured = json.loads(row["structured_json"]) if row["structured_json"] else {}
        raw_text = row["raw_text"] or ""
        structured_str = json.dumps(structured, ensure_ascii=False, indent=2)

        prompt = f"""你是一个严谨的植物科研实验助手。请审查以下从论文中提取的实验方案，找出其中可能存在的问题。

提取的实验方案：
{structured_str}

原始论文片段（用于交叉验证）：
{raw_text[:15000]}

请逐项检查：
1. **步骤完整性**：是否有遗漏的关键步骤？顺序是否正确？
2. **参数准确性**：温度、时间、转速、浓度等参数是否与原文一致？
3. **试剂信息**：试剂名称、浓度、用量是否正确？配制方法是否完整？
4. **仪器匹配**：仪器设备是否与实验方法匹配？
5. **逻辑一致性**：步骤之间是否有逻辑矛盾？
6. **可执行性**：一个新手看了能否直接操作？是否有模糊不清的描述？

返回JSON格式：
{{
  "issues": [
    {{
      "field": "steps/reagents/instruments/conditions/risks",
      "description": "具体问题描述",
      "suggestion": "修复建议"
    }}
  ]
}}

如果没有问题，返回 {{"issues": []}}

注意：
- 只报告**确实存在的问题**，不要吹毛求疵
- 问题描述要**具体**，指出是哪个步骤/试剂/参数有问题
- 修复建议要**可操作**
- 用中文返回"""

        result_text, usage = call_llm(prompt, api_settings)
        if not result_text:
            raise Exception("AI调用失败，请检查API Key和模型配置是否正确")

        result = _extract_json_from_llm(result_text)
        issues = result.get("issues", []) if result else []

        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='completed', issues_json=?, result_json=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (json.dumps(issues, ensure_ascii=False), json.dumps({"usage": usage}, ensure_ascii=False), task_id)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='failed', error_msg=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (str(e), task_id)
        )
        conn.commit()
        conn.close()


@app.route("/api/extract/fix", methods=["POST"])
@login_required
def api_extract_fix():
    """根据审查问题修复提取结果 - 后台异步执行"""
    data = request.json
    structured = data.get("structured", {})
    issues = data.get("issues", [])
    raw_text = data.get("raw_text", "")

    api_settings = get_user_api_settings()
    form_api_key = data.get("api_key", "")
    if form_api_key and not api_settings:
        api_settings = {"provider": "anthropic", "api_key": form_api_key, "model": "claude-sonnet-4-20250514", "base_url": ""}

    if not api_settings or not api_settings.get("api_key"):
        return jsonify({"error": "请先在个人中心配置API Key"})

    user_id = session["user_id"]
    conn = get_user_db()
    cur = conn.execute(
        "INSERT INTO extract_tasks (user_id, task_type, status, structured_json, issues_json, raw_text) VALUES (?, 'fix', 'pending', ?, ?, ?)",
        (user_id, json.dumps(structured, ensure_ascii=False), json.dumps(issues, ensure_ascii=False), raw_text[:50000])
    )
    task_id = cur.lastrowid
    conn.commit()
    conn.close()

    threading.Thread(
        target=_run_fix_task,
        args=(task_id, api_settings),
        daemon=True
    ).start()

    return jsonify({"task_id": task_id, "status": "pending"})


def _run_fix_task(task_id, api_settings):
    """后台执行修复任务"""
    try:
        conn = get_user_db()
        conn.execute("UPDATE extract_tasks SET status='processing' WHERE id=?", (task_id,))
        conn.commit()
        row = conn.execute("SELECT structured_json, issues_json, raw_text FROM extract_tasks WHERE id=?", (task_id,)).fetchone()
        conn.close()

        structured = json.loads(row["structured_json"]) if row["structured_json"] else {}
        issues = json.loads(row["issues_json"]) if row["issues_json"] else []
        raw_text = row["raw_text"] or ""

        structured_str = json.dumps(structured, ensure_ascii=False, indent=2)
        issues_str = json.dumps(issues, ensure_ascii=False, indent=2)

        prompt = f"""你是一个植物科研实验助手。请根据审查发现的问题，修正实验方案。

当前实验方案：
{structured_str}

需要修复的问题：
{issues_str}

原始论文片段（参考）：
{raw_text[:15000]}

请修正上述问题，返回**完整**的修正后JSON（格式与输入一致）：
{{
  "title": "...",
  "method_name": "...",
  "principle": "...",
  "steps": [{{"step": 1, "action": "...", "params": "...", "reagent": "...", "warning": "...", "source_confidence": "明确"}}],
  "reagents": [{{"name": "...", "concentration": "...", "amount": "...", "prep": "...", "source_confidence": "明确"}}],
  "instruments": ["..."],
  "conditions": {{"temperature": "...", "light": "...", "other": "..."}},
  "risks": [{{"point": "...", "solution": "...", "severity": "高/中/低"}}]
}}

注意：
- 只修正有问题的部分，没问题的保持不变
- 修正要基于原文，不要编造新内容
- 保留所有原始参数
- 返回完整JSON，不要只返回修改部分
- 用中文返回"""

        result_text, usage = call_llm(prompt, api_settings)
        if not result_text:
            raise Exception("AI调用失败，请检查API Key和模型配置是否正确")

        fixed = _extract_json_from_llm(result_text)
        if not fixed:
            raise Exception("AI返回格式异常，请重试。如反复失败可尝试更换模型。")

        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='completed', structured_json=?, result_json=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (json.dumps(fixed, ensure_ascii=False), json.dumps({"usage": usage}, ensure_ascii=False), task_id)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        conn = get_user_db()
        conn.execute(
            "UPDATE extract_tasks SET status='failed', error_msg=?, completed_at=CURRENT_TIMESTAMP WHERE id=?",
            (str(e), task_id)
        )
        conn.commit()
        conn.close()


@app.route("/api/extract/export", methods=["POST"])
def api_extract_export():
    """导出提取的实验方案为Word文档"""
    data = request.json
    structured = data.get("structured", {})
    fmt = data.get("format", "docx")

    doc = DocxDocument()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 设置Normal样式
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    method_name = structured.get("method_name", "") or structured.get("title", "")
    if method_name:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(6)
        run = title_p.add_run(method_name)
        set_run_font(run, 16, bold=True)

    # 来源
    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_p.paragraph_format.space_after = Pt(12)
    run = src_p.add_run("植研小白盒 · 文献实验提取")
    set_run_font(run, 9)
    run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    # 原理
    if structured.get("principle"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("实验原理：")
        set_run_font(run, 11, bold=True)
        run = p.add_run(structured["principle"])
        set_run_font(run, 11)

    # 实验步骤
    steps = structured.get("steps", [])
    if steps and isinstance(steps, list):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '2D9D78',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        run = p.add_run("实验步骤")
        set_run_font(run, 13, bold=True)
        run.font.color.rgb = RGBColor(0x2D, 0x9D, 0x78)

        for s in steps:
            step_num = s.get("step", "")
            action = s.get("action", "")
            params = s.get("params", "")
            warning = s.get("warning", "")
            reagent = s.get("reagent", "")
            confidence = s.get("source_confidence", "")

            # 步骤标题
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"步骤{step_num}：{action}")
            set_run_font(run, 10.5)

            # 参数标签
            if params:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.8)
                run = p.add_run(f"关键参数：{params}")
                set_run_font(run, 10)
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

            # 试剂
            if reagent:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.8)
                run = p.add_run(f"使用试剂：{reagent}")
                set_run_font(run, 10)
                run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

            # 注意事项
            if warning:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.8)
                run = p.add_run(f"⚠ {warning}")
                set_run_font(run, 10)
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    # 试剂配方
    reagents = structured.get("reagents", [])
    if reagents and isinstance(reagents, list):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '2D9D78',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        run = p.add_run("试剂配方")
        set_run_font(run, 13, bold=True)
        run.font.color.rgb = RGBColor(0x2D, 0x9D, 0x78)

        # 表格
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["试剂名称", "浓度", "用量", "配制方法"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            set_run_font(run, 10, bold=True)
            from docx.oxml import OxmlElement
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2C8')
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(shading)

        for r in reagents:
            row = table.add_row()
            vals = [r.get("name", ""), r.get("concentration", ""), r.get("amount", ""), r.get("prep", "")]
            for i, v in enumerate(vals):
                row.cells[i].text = ""
                run = row.cells[i].paragraphs[0].add_run(v or "-")
                set_run_font(run, 9)
                row.cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)

    # 仪器设备
    instruments = structured.get("instruments", [])
    if instruments and isinstance(instruments, list):
        add_paragraph(doc, "")
        p = doc.add_paragraph()
        run = p.add_run("仪器设备：")
        set_run_font(run, 11, bold=True)
        run = p.add_run("、".join(instruments))
        set_run_font(run, 11)

    # 实验条件
    conditions = structured.get("conditions", {})
    if conditions and isinstance(conditions, dict):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("实验条件：")
        set_run_font(run, 11, bold=True)
        cond_parts = []
        if conditions.get("temperature"):
            cond_parts.append(f"温度 {conditions['temperature']}")
        if conditions.get("light"):
            cond_parts.append(f"光照 {conditions['light']}")
        if conditions.get("other"):
            cond_parts.append(conditions["other"])
        run = p.add_run("；".join(cond_parts))
        set_run_font(run, 11)

    # 风险清单
    risks = structured.get("risks", [])
    if risks and isinstance(risks, list):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("风险清单")
        set_run_font(run, 13, bold=True)
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        for r in risks:
            severity = r.get("severity", "中")
            point = r.get("point", "")
            solution = r.get("solution", "")
            text = f"[{severity}] {point}"
            if solution:
                text += f" → {solution}"
            add_paragraph(doc, text, size=10)

    # 底部
    doc.add_paragraph()
    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(20)
    run = sign_p.add_run("实验日期：____________    操作人：____________    复核人：____________")
    set_run_font(run, 10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    title = method_name or "实验方案"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{title}.docx",
    )


# ========== 启动 ==========
if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("  植研小白盒 v2 启动中...")
    print("  访问地址: http://localhost:5000")
    print("=" * 50 + "\n")
    app.run(debug=True, port=5000)
